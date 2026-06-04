# -*- coding: utf-8 -*-
"""生成种子源模块控制器软件设计文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"f:\工作文档\2026\大功率可调恒流源\QT上位机\docs\种子源模块控制器_软件设计文档.docx"


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_with_header(doc, headers, rows, col_widths=None):
    """添加带表头样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, '2E75B6')

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_code_block(doc, code_text):
    """添加代码块样式的段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 添加底色
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F2F2F2')
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)

    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def set_heading_style(heading, font_size, color_rgb=None):
    """设置标题样式"""
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(font_size)
        if color_rgb:
            run.font.color.rgb = color_rgb


def build_document():
    doc = Document()

    # ============ 全局样式设置 ============
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.35

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ============ 封面标题 ============
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('种子源模块控制器')
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('软件设计文档')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    # 版本信息
    info_items = [
        ('项目名称', '种子源模块控制器（SeedSourceControl）'),
        ('版本', 'v1.0.1'),
        ('开发框架', 'Qt 6.5+ / C++11'),
        ('编译器', 'MSVC 2019/2022'),
        ('构建工具', 'CMake'),
        ('通信协议', '《种子源模块通信协议V1.3》'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{label}：')
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.name = '微软雅黑'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r2 = p.add_run(value)
        r2.font.size = Pt(12)
        r2.font.name = '微软雅黑'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # ============ 目录页 ============
    h = doc.add_heading('目录', level=1)
    set_heading_style(h, 18, RGBColor(0x1F, 0x49, 0x7D))

    toc_items = [
        '1. 项目概述',
        '2. 系统架构设计',
        '    2.1 分层架构',
        '    2.2 多线程架构',
        '    2.3 数据流',
        '3. 通信协议设计',
        '    3.1 帧格式',
        '    3.2 命令类型',
        '    3.3 寄存器地址',
        '4. 各子模块详细设计',
        '    4.1 协议解析器',
        '    4.2 寄存器管理器',
        '    4.3 命令系统',
        '    4.4 通信工作线程',
        '    4.5 设备数据模型',
        '    4.6 数据库管理器',
        '    4.7 报警管理器',
        '    4.8 配置管理器',
        '    4.9 UI面板',
        '5. 设计模式总结',
        '6. 线程安全设计',
        '7. 项目文件结构',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # ============ 1. 项目概述 ============
    h = doc.add_heading('1. 项目概述', level=1)
    set_heading_style(h, 18, RGBColor(0x1F, 0x49, 0x7D))

    add_table_with_header(doc,
        ['项目属性', '说明'],
        [
            ['项目名称', '种子源模块控制器（SeedSourceControl）'],
            ['版本', 'v1.0.1'],
            ['开发框架', 'Qt 6.5+ / C++11'],
            ['编译器', 'MSVC 2019/2022'],
            ['构建工具', 'CMake'],
            ['通信协议', '《种子源模块通信协议V1.3》'],
            ['功能概述', 'Windows上位机软件，用于控制和管理种子源模块设备，实现设备连接通信、参数配置、实时数据监控、数据可视化、报警管理、数据记录与导出等功能'],
        ],
        col_widths=[4, 12]
    )

    doc.add_paragraph()

    # ============ 2. 系统架构设计 ============
    h = doc.add_heading('2. 系统架构设计', level=1)
    set_heading_style(h, 18, RGBColor(0x1F, 0x49, 0x7D))

    # 2.1 分层架构
    h2 = doc.add_heading('2.1 分层架构', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['层级', '名称', '核心组件', '设计模式'],
        [
            ['1', '协议层（Protocol Layer）', '协议解析器、寄存器管理器', '策略模式、单例模式'],
            ['2', '通信层（Communication Layer）', '命令封装、通信工作线程、串口通信', '命令模式、生产者-消费者模式'],
            ['3', '数据模型层（Data Model Layer）', '设备数据模型、线程安全数据缓存', '观察者模式'],
            ['4', '可视化层（Visualization Layer）', '实时图表控件（QCustomPlot）', '—'],
            ['5', 'UI层（User Interface Layer）', '6个功能面板', '—'],
            ['6', '数据库层（Database Layer）', 'SQLite数据库管理', '单例模式'],
            ['7', '报警层（Alarm Layer）', '报警检测与触发', '单例模式'],
            ['8', '配置层（Config Layer）', '全局配置管理（QSettings持久化）', '单例模式'],
        ],
        col_widths=[1.5, 5, 5.5, 4]
    )

    doc.add_paragraph()

    # 2.2 多线程架构
    h2 = doc.add_heading('2.2 多线程架构', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['线程', '职责'],
        [
            ['主线程（UI线程）', '用户界面更新、用户输入处理、图表数据渲染'],
            ['通信线程', '串口数据收发、命令队列处理、超时重试机制'],
        ],
        col_widths=[4, 12]
    )

    doc.add_paragraph()
    p = doc.add_paragraph('线程间通信机制：')
    p.runs[0].bold = True
    bullets = [
        '通过Qt信号槽机制（QueuedConnection），确保线程安全',
        'QSerialPort在子线程中创建和管理，所有串口操作在同一线程执行',
        '主线程通过信号通知子线程进行连接/断开/发送命令操作',
    ]
    for b in bullets:
        bp = doc.add_paragraph(b, style='List Bullet')
        for run in bp.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 2.3 数据流
    h2 = doc.add_heading('2.3 数据流', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    flow_steps = [
        '用户操作 → UI面板发出信号 → MainWindow槽函数处理',
        'MainWindow创建命令对象（CommandFactory） → 通过信号传递给CommunicationWorker',
        'CommunicationWorker在子线程中执行命令 → 串口发送请求帧',
        '串口接收响应帧 → 协议解析 → 命令对象解析响应',
        '命令完成信号 → MainWindow更新数据模型 → UI面板更新显示',
    ]
    for i, step in enumerate(flow_steps, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f'Step {i}：')
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.name = '微软雅黑'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r2 = p.add_run(step)
        r2.font.size = Pt(10.5)
        r2.font.name = '微软雅黑'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # ============ 3. 通信协议设计 ============
    h = doc.add_heading('3. 通信协议设计', level=1)
    set_heading_style(h, 18, RGBColor(0x1F, 0x49, 0x7D))

    # 3.1 帧格式
    h2 = doc.add_heading('3.1 帧格式', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_code_block(doc, '[帧头0xAA][地址码][命令码][数据长度L][数据×L][校验和][帧尾0x55]')

    doc.add_paragraph()
    add_table_with_header(doc,
        ['字段', '长度（字节）', '说明'],
        [
            ['帧头', '1', '固定值0xAA'],
            ['地址码', '1', '设备地址，默认0x01'],
            ['命令码', '1', '命令类型编码'],
            ['数据长度L', '1', '数据段字节数'],
            ['数据', 'L', '命令参数数据'],
            ['校验和', '1', '从帧头到数据部分的累加和取低8位'],
            ['帧尾', '1', '固定值0x55'],
        ],
        col_widths=[3, 3, 10]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('命令超时：')
    r.bold = True
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2 = p.add_run('2000ms')
    r2.font.name = '微软雅黑'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 3.2 命令类型
    h2 = doc.add_heading('3.2 命令类型', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['命令码', '命令名称', '参数说明', '返回数据'],
        [
            ['0x01', '读寄存器', '基地址 + 偏移地址', '寄存器值'],
            ['0x02', '写寄存器', '基地址 + 偏移地址 + 4字节数值（大端序）', '写操作结果'],
            ['0x03', '读状态', '无参数', '12字节数据（电流4 + 温度4 + 功率4）'],
            ['0x04', '设备控制', '控制码（0x01启动/0x02停止/0x03重置/0x04校准）', '控制操作结果'],
        ],
        col_widths=[2, 3, 6, 5]
    )

    doc.add_paragraph()

    # 3.3 寄存器地址
    h2 = doc.add_heading('3.3 寄存器地址', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    p = doc.add_paragraph('地址计算公式：')
    p.runs[0].bold = True
    add_code_block(doc, '地址 = 基地址 << 8 | 偏移地址')

    doc.add_paragraph()
    add_table_with_header(doc,
        ['地址', '名称', '说明'],
        [
            ['0x00', 'STATUS', '设备状态寄存器'],
            ['0x01', 'ALERT', '报警状态寄存器'],
            ['0x02', 'CUR', '电流控制寄存器（0-1000mA）'],
            ['0x03', 'TEMP', '温度控制寄存器（-40~125°C）'],
            ['0x04', 'POWER', '功率控制寄存器（0-10000mW）'],
            ['0x05', 'CONFIG', '配置寄存器'],
            ['0x06', 'SYSTEM', '系统控制寄存器'],
            ['0x07', 'DEVINFO', '设备信息寄存器'],
        ],
        col_widths=[2.5, 3, 10.5]
    )

    doc.add_page_break()

    # ============ 4. 各子模块详细设计 ============
    h = doc.add_heading('4. 各子模块详细设计', level=1)
    set_heading_style(h, 18, RGBColor(0x1F, 0x49, 0x7D))

    # 4.1 协议解析器
    h2 = doc.add_heading('4.1 协议解析器（SeedSourceProtocolParser）', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['属性', '说明'],
        [
            ['继承', 'QObject + IProtocolParser接口'],
            ['设计模式', '策略模式（IProtocolParser接口允许替换不同协议实现）'],
            ['设备地址', '默认0x01'],
        ],
        col_widths=[3, 13]
    )
    doc.add_paragraph()
    p = doc.add_paragraph('核心方法：')
    p.runs[0].bold = True
    methods = [
        'buildFrame() — 构建协议帧',
        'parseFrame() — 解析协议帧',
        'calculateChecksum() — 计算校验和',
    ]
    for m in methods:
        add_code_block(doc, m)

    # 4.2 寄存器管理器
    h2 = doc.add_heading('4.2 寄存器管理器（RegisterManager）', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['属性', '说明'],
        [
            ['设计模式', '单例模式（双重检查锁定）'],
            ['线程安全', 'QMutex保护所有寄存器操作'],
            ['功能', '寄存器定义管理、读写操作、值范围检查'],
        ],
        col_widths=[3, 13]
    )
    doc.add_paragraph()
    p = doc.add_paragraph('寄存器定义结构：')
    p.runs[0].bold = True
    add_code_block(doc, '名称、基地址、偏移、大小、默认值、最小值、最大值、单位、描述')

    # 4.3 命令系统
    h2 = doc.add_heading('4.3 命令系统（ICommand及其子类）', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['组件', '说明'],
        [
            ['设计模式', '命令模式'],
            ['基类ICommand', '定义命令接口（execute/onResponse/onTimeout）'],
            ['ReadRegisterCommand', '读寄存器命令'],
            ['WriteRegisterCommand', '写寄存器命令'],
            ['ReadStatusCommand', '读状态命令'],
            ['ControlDeviceCommand', '设备控制命令'],
            ['CommandFactory', '工厂模式创建命令对象'],
            ['超时处理', 'QTimer，超时时间2000ms'],
            ['线程安全', '使用QMetaObject::invokeMethod确保跨线程调用安全'],
        ],
        col_widths=[4, 12]
    )

    # 4.4 通信工作线程
    h2 = doc.add_heading('4.4 通信工作线程（CommunicationWorker）', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['属性', '说明'],
        [
            ['继承', 'QThread'],
            ['设计模式', '生产者-消费者模式（命令队列）'],
        ],
        col_widths=[3, 13]
    )

    doc.add_paragraph()
    p = doc.add_paragraph('线程安全措施：')
    p.runs[0].bold = True
    safety_items = [
        'm_stateMutex保护m_connected和m_state',
        'm_queueMutex保护命令队列',
        'm_pendingMutex保护待处理命令映射',
        'm_bufferMutex保护接收缓冲区',
        'QSerialPort在子线程中创建和管理',
        '主线程通过信号（connectRequested/disconnectRequested/sendCommandRequested）通知子线程',
    ]
    for item in safety_items:
        bp = doc.add_paragraph(item, style='List Bullet')
        for run in bp.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    p = doc.add_paragraph('关键功能：')
    p.runs[0].bold = True
    funcs = [
        'findCompleteFrame() — 在接收缓冲区中查找完整帧',
        '5秒定时器检测串口连接状态',
    ]
    for f in funcs:
        bp = doc.add_paragraph(f, style='List Bullet')
        for run in bp.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 4.5 设备数据模型
    h2 = doc.add_heading('4.5 设备数据模型（DeviceDataModel）', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['属性', '说明'],
        [
            ['设计模式', '观察者模式（dataUpdated/alarmTriggered信号）'],
            ['线程安全', 'QReadWriteLock保护实时数据；DataCache使用独立的QReadWriteLock'],
            ['数据缓存', '环形缓冲区，默认最大10000条记录'],
            ['报警检测', 'checkAlarms()方法检测位域报警码'],
        ],
        col_widths=[3, 13]
    )

    # 4.6 数据库管理器
    h2 = doc.add_heading('4.6 数据库管理器（DatabaseManager）', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['属性', '说明'],
        [
            ['设计模式', '单例模式'],
            ['数据库', 'SQLite（QSQLITE驱动）'],
            ['功能', '数据存储、查询、报警记录、CSV导出'],
        ],
        col_widths=[3, 13]
    )

    doc.add_paragraph()
    p = doc.add_paragraph('表结构：')
    p.runs[0].bold = True

    add_table_with_header(doc,
        ['表名', '字段'],
        [
            ['data', '时间戳、电流、温度、功率、状态、报警'],
            ['alarms', '时间戳、报警码、消息'],
        ],
        col_widths=[3, 13]
    )
    doc.add_paragraph()
    p = doc.add_paragraph('索引：')
    p.runs[0].bold = True
    add_code_block(doc, 'idx_data_timestamp')

    # 4.7 报警管理器
    h2 = doc.add_heading('4.7 报警管理器（AlarmManager）', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['属性', '说明'],
        [
            ['设计模式', '单例模式'],
            ['线程安全', 'QReadWriteLock保护阈值数据'],
            ['阈值配置', '通过alarm_config.ini文件持久化'],
        ],
        col_widths=[3, 13]
    )

    doc.add_paragraph()
    p = doc.add_paragraph('报警类型：')
    p.runs[0].bold = True

    add_table_with_header(doc,
        ['报警码', '报警类型', '说明'],
        [
            ['0x01', '过流', '电流超过阈值'],
            ['0x02', '过温', '温度超过阈值'],
            ['0x04', '过功率', '功率超过阈值'],
            ['0x08', '电压错误', '电压异常'],
            ['0x10', '通信错误', '通信超时或校验失败'],
            ['0x20', '硬件故障', '硬件异常'],
        ],
        col_widths=[2.5, 3, 10.5]
    )

    # 4.8 配置管理器
    h2 = doc.add_heading('4.8 配置管理器（ConfigManager）', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['属性', '说明'],
        [
            ['设计模式', '单例模式'],
            ['线程安全', 'QRecursiveMutex保护所有配置成员和QSettings操作'],
            ['持久化', 'QSettings（INI格式）'],
        ],
        col_widths=[3, 13]
    )

    doc.add_paragraph()
    p = doc.add_paragraph('配置项：')
    p.runs[0].bold = True
    config_items = ['串口参数', '轮询间隔', '超时时间', '日志路径', '数据库路径']
    for item in config_items:
        bp = doc.add_paragraph(item, style='List Bullet')
        for run in bp.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 4.9 UI面板
    h2 = doc.add_heading('4.9 UI面板', level=2)
    set_heading_style(h2, 14, RGBColor(0x2E, 0x75, 0xB6))

    add_table_with_header(doc,
        ['面板名称', '功能描述'],
        [
            ['ConnectionPanel', '串口选择、参数配置、连接/断开按钮'],
            ['ControlPanel', '启动/停止/重置/校准按钮、电流/温度/功率参数设置'],
            ['StatusPanel', '设备状态、实时电流/温度/功率/报警显示'],
            ['RealtimePlotPanel', '电流/温度/功率实时曲线（基于QCustomPlot）'],
            ['AlarmPanel', '报警历史记录表格'],
            ['LogPanel', '通信/操作/错误日志显示'],
        ],
        col_widths=[4, 12]
    )

    doc.add_page_break()

    # ============ 5. 设计模式总结 ============
    h = doc.add_heading('5. 设计模式总结', level=1)
    set_heading_style(h, 18, RGBColor(0x1F, 0x49, 0x7D))

    add_table_with_header(doc,
        ['序号', '设计模式', '应用场景', '说明'],
        [
            ['1', '命令模式', 'ICommand封装设备操作', '支持撤销和超时'],
            ['2', '观察者模式', 'DeviceDataModel发出数据更新信号', 'UI面板自动更新'],
            ['3', '策略模式', 'IProtocolParser接口', '可替换不同协议实现'],
            ['4', '工厂模式', 'CommandFactory创建命令对象', '创建各种命令对象'],
            ['5', '单例模式', 'RegisterManager/DatabaseManager/\nAlarmManager/ConfigManager', '全局唯一实例'],
            ['6', '生产者-消费者模式', 'CommunicationWorker的命令队列', '命令异步处理'],
        ],
        col_widths=[1.5, 3.5, 5.5, 5.5]
    )

    doc.add_paragraph()

    # ============ 6. 线程安全设计 ============
    h = doc.add_heading('6. 线程安全设计', level=1)
    set_heading_style(h, 18, RGBColor(0x1F, 0x49, 0x7D))

    safety_design = [
        '所有共享资源使用QMutex/QReadWriteLock保护',
        'QSerialPort在子线程中创建和管理，避免跨线程访问',
        'QTimer在所属线程中创建和销毁',
        'ICommand的onResponse通过QMetaObject::invokeMethod跨线程调用',
        '主线程与子线程通过Qt信号槽（QueuedConnection）通信',
    ]
    for i, item in enumerate(safety_design, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f'{i}. ')
        r1.bold = True
        r1.font.name = '微软雅黑'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r2 = p.add_run(item)
        r2.font.name = '微软雅黑'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # ============ 7. 项目文件结构 ============
    h = doc.add_heading('7. 项目文件结构', level=1)
    set_heading_style(h, 18, RGBColor(0x1F, 0x49, 0x7D))

    file_structure = [
        # (文件路径, 职责说明)
        # 根目录
        ('CMakeLists.txt', 'CMake构建配置文件，定义项目结构、源文件、依赖库'),
        ('resources.qrc', 'Qt资源文件，管理图标、样式表等资源'),
        ('styles/default.qss', '全局样式表文件，定义UI控件外观风格'),
        # src根
        ('src/main.cpp', '应用程序入口，初始化Qt应用、加载配置、启动主窗口'),
        ('src/mainwindow.h', '主窗口头文件，声明主窗口类及各面板管理接口'),
        ('src/mainwindow.cpp', '主窗口实现，协调各子模块、处理信号槽连接、面板布局'),
        # protocol
        ('src/protocol/protocolparser.h', '协议解析器头文件，声明IProtocolParser接口和SeedSourceProtocolParser类'),
        ('src/protocol/protocolparser.cpp', '协议解析器实现，帧构建、帧解析、校验和计算'),
        ('src/protocol/registermanager.h', '寄存器管理器头文件，声明RegisterManager单例类及寄存器定义结构'),
        ('src/protocol/registermanager.cpp', '寄存器管理器实现，寄存器CRUD操作、值范围检查、线程安全保护'),
        # communication
        ('src/communication/command.h', '命令系统头文件，声明ICommand基类及各命令子类、CommandFactory'),
        ('src/communication/command.cpp', '命令系统实现，各命令的execute/onResponse/onTimeout逻辑'),
        ('src/communication/communicationworker.h', '通信工作线程头文件，声明CommunicationWorker类'),
        ('src/communication/communicationworker.cpp', '通信工作线程实现，串口管理、命令队列处理、帧解析、超时重试'),
        # model
        ('src/model/devicedatamodel.h', '设备数据模型头文件，声明DeviceDataModel类及DataCache结构'),
        ('src/model/devicedatamodel.cpp', '设备数据模型实现，实时数据管理、数据缓存、报警检测'),
        # database
        ('src/database/databasemanager.h', '数据库管理器头文件，声明DatabaseManager单例类'),
        ('src/database/databasemanager.cpp', '数据库管理器实现，SQLite建表、数据存储查询、CSV导出'),
        # alarm
        ('src/alarm/alarmmanager.h', '报警管理器头文件，声明AlarmManager单例类及报警类型枚举'),
        ('src/alarm/alarmmanager.cpp', '报警管理器实现，报警检测、阈值管理、报警信号触发'),
        # config
        ('src/config/configmanager.h', '配置管理器头文件，声明ConfigManager单例类'),
        ('src/config/configmanager.cpp', '配置管理器实现，QSettings读写、全局配置项管理'),
        # visualization
        ('src/visualization/realtimeplotwidget.h', '实时图表控件头文件，声明RealtimePlotWidget类'),
        ('src/visualization/realtimeplotwidget.cpp', '实时图表控件实现，QCustomPlot封装、曲线绘制、数据更新'),
        # ui
        ('src/ui/connectionpanel.h', '连接面板头文件，声明ConnectionPanel类'),
        ('src/ui/connectionpanel.cpp', '连接面板实现，串口选择、参数配置、连接/断开操作'),
        ('src/ui/controlpanel.h', '控制面板头文件，声明ControlPanel类'),
        ('src/ui/controlpanel.cpp', '控制面板实现，设备启停控制、电流/温度/功率参数设置'),
        ('src/ui/statuspanel.h', '状态面板头文件，声明StatusPanel类'),
        ('src/ui/statuspanel.cpp', '状态面板实现，设备状态、实时电流/温度/功率/报警显示'),
        ('src/ui/realtimeplotpanel.h', '实时曲线面板头文件，声明RealtimePlotPanel类'),
        ('src/ui/realtimeplotpanel.cpp', '实时曲线面板实现，集成RealtimePlotWidget、管理曲线数据源'),
        ('src/ui/alarmpanel.h', '报警面板头文件，声明AlarmPanel类'),
        ('src/ui/alarmpanel.cpp', '报警面板实现，报警历史记录表格显示与管理'),
        ('src/ui/logpanel.h', '日志面板头文件，声明LogPanel类'),
        ('src/ui/logpanel.cpp', '日志面板实现，通信/操作/错误日志显示与过滤'),
        ('src/ui/datatablepanel.h', '数据表格面板头文件，声明DataTablePanel类'),
        ('src/ui/datatablepanel.cpp', '数据表格面板实现，历史数据表格展示与导出'),
    ]

    add_table_with_header(doc,
        ['文件路径', '职责说明'],
        [list(row) for row in file_structure],
        col_widths=[6, 10]
    )

    # ============ 保存文档 ============
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f'文档已成功生成：{OUTPUT_PATH}')


if __name__ == '__main__':
    build_document()

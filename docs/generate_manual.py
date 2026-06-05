# -*- coding: utf-8 -*-
"""生成种子源模块控制器使用说明书 Word 文档（v1.0.1 更新版）"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"f:\工作文档\2026\大功率可调恒流源\QT上位机\docs\种子源模块控制器_使用说明书.docx"

# ── 颜色常量 ──
BLUE_HEADER = "1F497D"
BLUE_LIGHT  = "D6E4F0"
WHITE       = "FFFFFF"


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, size=Pt(11), bold=False, color=None):
    """统一设置 run 的中文字体"""
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level):
    """添加标题并设置中文字体"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=run.font.size, bold=True)
    return heading


def add_para(doc, text, bold=False, indent=False, space_after=Pt(4)):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_bullet(doc, text, level=0):
    """添加列表项"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.75)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    p.clear()
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_numbered(doc, text, level=0):
    """添加有序列表项"""
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.75)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    p.clear()
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加带蓝色表头和交替行底色的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, BLUE_HEADER)

    # 数据行
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            set_run_font(run)
            if i % 2 == 1:
                set_cell_shading(cell, BLUE_LIGHT)

    # 列宽
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)

    return table


def build_doc():
    doc = Document()

    # ── 全局默认字体 ──
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ── 页边距 ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ═══════════════════════════════════════════
    # 封面
    # ═══════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("种子源模块控制器")
    set_run_font(title_run, size=Pt(28), bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("使 用 说 明 书")
    set_run_font(sub_run, size=Pt(22), color=RGBColor(0x1F, 0x49, 0x7D))

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.paragraph_format.space_before = Pt(20)
    ver_run = ver_p.add_run("版本 v1.0.1")
    set_run_font(ver_run, size=Pt(14), color=RGBColor(0x59, 0x56, 0x59))

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 1. 软件简介
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "1. 软件简介", level=1)
    add_bullet(doc, "软件名称：种子源模块控制器（SeedSourceControl）")
    add_bullet(doc, "版本：v1.0.1")
    add_bullet(doc, "运行环境：Windows 7/10/11（64位）")
    add_bullet(doc, "软件用途：用于控制和管理种子源模块设备，通过串口与设备通信")
    add_bullet(doc, "通信协议：《种子源模块通信协议V1.3》")
    add_bullet(doc, "主要功能：设备连接通信、参数配置、实时数据监控、数据可视化、报警管理、数据记录与导出")

    # ═══════════════════════════════════════════
    # 2. 安装与启动
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "2. 安装与启动", level=1)

    add_heading_styled(doc, "2.1 系统要求", level=2)
    add_bullet(doc, "操作系统：Windows 7/10/11 64位")
    add_bullet(doc, "无需安装Qt环境（已集成运行时依赖）")
    add_bullet(doc, "至少一个可用串口（物理串口或虚拟串口）")

    add_heading_styled(doc, "2.2 启动方式", level=2)
    add_bullet(doc, "双击 SeedSourceControl.exe 即可启动")
    add_bullet(doc, "程序位于 build\\deploy\\Debug\\ 目录")
    add_bullet(doc, "程序以Windows GUI模式运行，不会弹出命令行窗口")

    add_heading_styled(doc, "2.3 目录说明", level=2)
    add_bullet(doc, "SeedSourceControl.exe — 主程序")
    add_bullet(doc, "platforms\\ — Qt平台插件（qwindowsd.dll等）")
    add_bullet(doc, "sqldrivers\\ — SQLite数据库驱动（qsqlited.dll）")
    add_bullet(doc, "imageformats\\ — 图像格式插件")
    add_bullet(doc, "styles\\ — Qt样式插件")

    # ═══════════════════════════════════════════
    # 3. 界面说明
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "3. 界面说明", level=1)

    add_heading_styled(doc, "3.1 主界面布局", level=2)
    add_para(doc, "主界面分为左右两个区域，通过可拖拽的分割条调整大小：")
    add_bullet(doc, "左侧区域（约30%宽度，从上到下）：")
    add_bullet(doc, "1. 连接面板（Serial Port Configuration）", level=1)
    add_bullet(doc, "2. 控制面板（Device Control + Parameter Settings）", level=1)
    add_bullet(doc, "3. 状态面板（Device Status）", level=1)
    add_bullet(doc, "右侧区域（约70%宽度，从上到下）：")
    add_bullet(doc, "1. 实时图表面板", level=1)
    add_bullet(doc, "2. 报警面板", level=1)
    add_bullet(doc, "3. 日志面板", level=1)

    add_heading_styled(doc, "3.2 连接面板（Serial Port Configuration）", level=2)
    add_bullet(doc, "Port：串口选择下拉框，显示系统可用COM端口")
    add_bullet(doc, "Baud Rate：波特率下拉框，默认115200，支持所有标准波特率")
    add_bullet(doc, "Data Bits：数据位下拉框，选项5/6/7/8，默认8")
    add_bullet(doc, "Parity：校验位下拉框，选项None/Even/Odd/Mark/Space，默认None")
    add_bullet(doc, "Stop Bits：停止位下拉框，选项1/1.5/2，默认1")
    add_bullet(doc, "Flow Control：流控制下拉框，选项None/Hardware/Software，默认None")
    add_bullet(doc, "Refresh按钮：刷新可用串口列表")
    add_bullet(doc, "Connect/Disconnect按钮：连接/断开设备（连接后按钮文字切换为Disconnect）")
    add_bullet(doc, "Status标签：显示连接状态（绿色Connected/红色Disconnected）")
    add_bullet(doc, "注意：连接后所有配置控件将被禁用，需断开后才能修改")

    add_heading_styled(doc, "3.3 控制面板", level=2)
    add_para(doc, "设备控制区（Device Control）：", bold=True)
    add_bullet(doc, "Start按钮（绿色）：启动设备运行")
    add_bullet(doc, "Stop按钮（红色）：停止设备运行")
    add_bullet(doc, "Reset按钮（橙色）：重置设备")
    add_bullet(doc, "Calibrate按钮（蓝色）：执行设备校准")
    add_para(doc, "参数设置区（Parameter Settings）：", bold=True)
    add_bullet(doc, "Target Current：目标电流，范围0-1000.00 mA，步进0.01")
    add_bullet(doc, "Target Temperature：目标温度，范围-40.00-125.00 °C，步进0.01")
    add_bullet(doc, "Target Power：目标功率，范围0-10000.00 mW，步进0.01")
    add_bullet(doc, "注意：参数变更会实时发送到设备（需设备已连接且运行中）")

    add_heading_styled(doc, "3.4 状态面板（Device Status）", level=2)
    add_bullet(doc, "Status：设备状态（Idle灰色/Running绿色/Error红色）")
    add_bullet(doc, "Current：实时电流值（mA），格式xx.xx mA")
    add_bullet(doc, "Temperature：实时温度值（°C），>80°C红色，>60°C橙色")
    add_bullet(doc, "Power：实时功率值（W），格式xx.xx W")
    add_bullet(doc, "Alarm：报警状态（No Alarm绿色/具体报警红色）")
    add_bullet(doc, "报警类型：Over Current/Over Temp/Over Power/Voltage Error/Comm Error/HW Fault", level=1)

    add_heading_styled(doc, "3.5 实时图表面板", level=2)
    add_bullet(doc, "电流曲线：实时显示电流变化趋势")
    add_bullet(doc, "温度曲线：实时显示温度变化趋势")
    add_bullet(doc, "功率曲线：实时显示功率变化趋势")
    add_bullet(doc, "支持鼠标缩放和平移操作")

    add_heading_styled(doc, "3.6 报警面板", level=2)
    add_bullet(doc, "报警列表：显示所有报警记录")
    add_bullet(doc, "报警信息包括：时间、报警码、报警描述")
    add_bullet(doc, "清除按钮：清除所有报警记录")

    add_heading_styled(doc, "3.7 日志面板", level=2)
    add_bullet(doc, "日志级别：Info（信息）、Warning（警告）、Error（错误）")
    add_bullet(doc, "显示内容：通信日志（TX/RX十六进制数据）、操作日志、错误日志")
    add_bullet(doc, "清除按钮：清除所有日志")
    add_bullet(doc, "保存按钮：保存日志到文件")
    add_bullet(doc, "自动滚动到最新日志")

    add_heading_styled(doc, "3.8 菜单栏", level=2)
    add_bullet(doc, "文件(&F)：退出(&X) — Ctrl+Q")
    add_bullet(doc, "视图(&V)：视图选项")
    add_bullet(doc, "工具(&T)：工具选项")
    add_bullet(doc, "帮助(&H)：关于(&A) — 显示\"种子源模块控制器 v1.0 基于Qt框架开发\"")

    add_heading_styled(doc, "3.9 工具栏", level=2)
    add_bullet(doc, "连接按钮：快速连接设备")
    add_bullet(doc, "断开按钮：快速断开设备")
    add_bullet(doc, "启动按钮：快速启动设备")
    add_bullet(doc, "停止按钮：快速停止设备")

    add_heading_styled(doc, "3.10 状态栏", level=2)
    add_bullet(doc, "显示当前操作状态：就绪/已连接/已断开/运行中.../已停止")

    # ═══════════════════════════════════════════
    # 4. 操作流程
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "4. 操作流程", level=1)

    add_heading_styled(doc, "4.1 连接设备", level=2)
    add_numbered(doc, "在连接面板的Port下拉框选择正确的串口（如COM3）")
    add_numbered(doc, "设置串口参数（波特率等，需与设备一致，默认115200-8-N-1）")
    add_numbered(doc, "点击\"Connect\"按钮")
    add_numbered(doc, "观察Status标签变为绿色\"Connected\"")
    add_numbered(doc, "状态栏显示\"已连接\"")
    add_numbered(doc, "日志面板显示\"Device connected: COMx\"")

    add_heading_styled(doc, "4.2 启动设备", level=2)
    add_numbered(doc, "确认设备已连接（Status标签显示Connected）")
    add_numbered(doc, "在控制面板点击\"Start\"按钮（绿色）")
    add_numbered(doc, "状态栏显示\"运行中...\"")
    add_numbered(doc, "日志面板显示\"设备启动\"")
    add_numbered(doc, "实时图表开始显示数据曲线")
    add_numbered(doc, "状态面板开始更新实时数据")

    add_heading_styled(doc, "4.3 设置参数", level=2)
    add_numbered(doc, "确认设备已连接且正在运行")
    add_numbered(doc, "在控制面板调整参数：")
    add_numbered(doc, "Target Current：设置目标电流（0-1000 mA）", level=1)
    add_numbered(doc, "Target Temperature：设置目标温度（-40~125 °C）", level=1)
    add_numbered(doc, "Target Power：设置目标功率（0-10000 mW）", level=1)
    add_numbered(doc, "参数变更会自动发送写寄存器命令到设备")
    add_numbered(doc, "日志面板显示\"设置目标电流/温度/功率: xxx\"")
    add_numbered(doc, "观察状态面板和图表的实时变化")

    add_heading_styled(doc, "4.4 监控报警", level=2)
    add_numbered(doc, "当设备出现异常时，报警面板自动显示报警信息")
    add_numbered(doc, "状态面板的Alarm标签变为红色并显示具体报警")
    add_numbered(doc, "报警类型：")
    add_bullet(doc, "0x01 Over Current：过流", level=1)
    add_bullet(doc, "0x02 Over Temperature：过温", level=1)
    add_bullet(doc, "0x04 Over Power：过功率", level=1)
    add_bullet(doc, "0x08 Voltage Error：电压错误", level=1)
    add_bullet(doc, "0x10 Communication Error：通信错误", level=1)
    add_bullet(doc, "0x20 Hardware Fault：硬件故障", level=1)
    add_numbered(doc, "日志面板同时记录报警详情（Error级别）")

    add_heading_styled(doc, "4.5 停止和断开", level=2)
    add_numbered(doc, "点击\"Stop\"按钮（红色）停止设备运行")
    add_numbered(doc, "状态栏显示\"已停止\"")
    add_numbered(doc, "点击\"Disconnect\"按钮断开串口连接")
    add_numbered(doc, "Status标签变为红色\"Disconnected\"")
    add_numbered(doc, "所有配置控件恢复可编辑状态")
    add_numbered(doc, "数据已自动保存到SQLite数据库")

    # ═══════════════════════════════════════════
    # 5. 数据管理
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "5. 数据管理", level=1)

    add_heading_styled(doc, "5.1 数据存储", level=2)
    add_bullet(doc, "所有实时数据自动保存到SQLite数据库")
    add_bullet(doc, "数据库文件位置：用户AppData目录")
    add_bullet(doc, "数据表字段：时间戳、电流(mA)、温度(°C)、功率(mW)、状态、报警码")
    add_bullet(doc, "报警表字段：时间戳、报警码、报警消息")
    add_bullet(doc, "数据库索引：idx_data_timestamp")

    add_heading_styled(doc, "5.2 数据导出", level=2)
    add_bullet(doc, "支持CSV格式导出")
    add_bullet(doc, "导出内容：Timestamp, Current, Temperature, Power, Status, Alarm")
    add_bullet(doc, "按时间戳排序")

    add_heading_styled(doc, "5.3 数据缓存", level=2)
    add_bullet(doc, "内存中维护最近10000条数据的环形缓冲区")
    add_bullet(doc, "支持按时间范围查询历史数据")

    # ═══════════════════════════════════════════
    # 6. 配置说明
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "6. 配置说明", level=1)

    add_heading_styled(doc, "6.1 串口配置", level=2)
    add_bullet(doc, "串口参数通过界面设置")
    add_bullet(doc, "配置自动保存到config.ini，下次启动自动加载")
    add_bullet(doc, "默认配置：115200-8-N-1-NoFlowControl")

    add_heading_styled(doc, "6.2 报警阈值配置", level=2)
    add_bullet(doc, "配置文件：alarm_config.ini（位于日志文件同目录）")
    add_bullet(doc, "可配置项：")
    add_bullet(doc, "Current/Min、Current/Max、Current/Enabled", level=1)
    add_bullet(doc, "Temperature/Min、Temperature/Max、Temperature/Enabled", level=1)
    add_bullet(doc, "Power/Min、Power/Max、Power/Enabled", level=1)
    add_bullet(doc, "默认阈值：电流0-1000mA、温度-40-85°C、功率0-10000mW")
    add_bullet(doc, "默认状态：所有报警检测禁用（Enabled=false）")
    add_bullet(doc, "修改后重启程序生效")

    add_heading_styled(doc, "6.3 应用配置", level=2)
    add_bullet(doc, "配置文件：config.ini（位于QStandardPaths::AppConfigLocation）")
    add_bullet(doc, "配置项：")
    add_bullet(doc, "Serial/Port、Serial/BaudRate、Serial/DataBits、Serial/Parity、Serial/StopBits、Serial/FlowControl", level=1)
    add_bullet(doc, "General/PollInterval（默认100ms）", level=1)
    add_bullet(doc, "General/Timeout（默认2000ms）", level=1)
    add_bullet(doc, "Paths/LogFile、Paths/Database", level=1)

    # ═══════════════════════════════════════════
    # 7. 通信协议参考
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "7. 通信协议参考", level=1)

    add_heading_styled(doc, "7.1 帧格式", level=2)
    add_para(doc, "[帧头0xAA][地址码][命令码][数据长度L][数据×L][校验和][帧尾0x55]")

    add_heading_styled(doc, "7.2 命令码", level=2)
    add_styled_table(
        doc,
        ["命令码", "名称", "发送参数", "响应数据"],
        [
            ["0x01", "读寄存器", "baseAddr(1B)+offset(1B)", "4字节大端序值"],
            ["0x02", "写寄存器", "baseAddr(1B)+offset(1B)+value(4B)", "确认"],
            ["0x03", "读状态", "无", "12字节(current+temp+power)"],
            ["0x04", "设备控制", "actionCode(1B)", "确认"],
        ],
        col_widths=[2.5, 2.5, 5.0, 4.0],
    )

    add_heading_styled(doc, "7.3 控制码", level=2)
    add_styled_table(
        doc,
        ["控制码", "操作"],
        [
            ["0x01", "启动"],
            ["0x02", "停止"],
            ["0x03", "重置"],
            ["0x04", "校准"],
        ],
        col_widths=[4.0, 5.0],
    )

    # ═══════════════════════════════════════════
    # 8. 常见问题
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "8. 常见问题", level=1)

    add_heading_styled(doc, "Q1: 无法找到串口", level=3)
    add_bullet(doc, "检查设备是否已连接到电脑")
    add_bullet(doc, "检查USB转串口驱动是否已安装")
    add_bullet(doc, "点击\"Refresh\"按钮重新扫描可用端口")
    add_bullet(doc, "检查设备管理器中串口是否正常识别")

    add_heading_styled(doc, "Q2: 连接失败", level=3)
    add_bullet(doc, "确认串口参数与设备一致（默认115200-8-N-1）")
    add_bullet(doc, "确认串口未被其他程序占用")
    add_bullet(doc, "检查设备是否已上电")
    add_bullet(doc, "查看日志面板的错误信息")

    add_heading_styled(doc, "Q3: 数据不更新", level=3)
    add_bullet(doc, "确认设备已连接且已启动（点击Start按钮）")
    add_bullet(doc, "检查串口通信是否正常（查看日志面板TX/RX数据）")
    add_bullet(doc, "确认轮询定时器正在运行")
    add_bullet(doc, "检查波特率设置是否正确")

    add_heading_styled(doc, "Q4: 参数设置无效", level=3)
    add_bullet(doc, "确认设备已连接且正在运行")
    add_bullet(doc, "参数变更只在设备运行状态下生效")
    add_bullet(doc, "检查参数值是否在有效范围内")

    add_heading_styled(doc, "Q5: 报警阈值如何修改", level=3)
    add_bullet(doc, "通过alarm_config.ini文件修改")
    add_bullet(doc, "修改后重启程序生效")
    add_bullet(doc, "默认所有报警检测为禁用状态")

    add_heading_styled(doc, "Q6: 如何导出数据", level=3)
    add_bullet(doc, "使用数据表格面板的导出功能")
    add_bullet(doc, "导出为CSV格式，可用Excel打开")

    # ═══════════════════════════════════════════
    # 9. 技术参数
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "9. 技术参数", level=1)

    add_styled_table(
        doc,
        ["参数", "范围", "默认值", "说明"],
        [
            ["电流", "0-1000 mA", "0 mA", "电流控制范围"],
            ["温度", "-40~125 °C", "25 °C", "温度控制范围"],
            ["功率", "0-10000 mW", "0 mW", "功率控制范围"],
            ["波特率", "9600-115200", "115200", "串口通信速率"],
            ["轮询间隔", "—", "100 ms", "设备状态轮询周期"],
            ["命令超时", "—", "2000 ms", "命令响应超时时间"],
            ["数据缓存", "—", "10000条", "内存数据缓存大小"],
            ["帧缓冲区", "—", "4096字节", "接收缓冲区上限"],
        ],
        col_widths=[3.0, 3.5, 3.0, 4.5],
    )

    # ═══════════════════════════════════════════
    # 10. 联系方式
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "10. 联系方式", level=1)
    add_para(doc, "如有问题或建议，请联系开发团队。")

    # ── 保存 ──
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()

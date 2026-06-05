# -*- coding: utf-8 -*-
"""生成种子源模块控制器软件设计文档"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 路径常量 ──
BASE_DIR = r'f:\工作文档\2026\大功率可调恒流源\QT上位机'
IMAGE_DIR = os.path.join(BASE_DIR, 'docs', 'images')
OUTPUT_PATH = os.path.join(BASE_DIR, 'docs', '种子源模块控制器_软件设计文档.docx')

# ── 颜色常量 ──
BLUE_TITLE = RGBColor(0x1F, 0x4E, 0x79)
GRAY_CAPTION = RGBColor(0x80, 0x80, 0x80)
BLACK = RGBColor(0, 0, 0)

# ── 字体常量 ──
FONT_NAME = '微软雅黑'
FONT_NAME_ASCII = 'Microsoft YaHei'

# 中文引号常量，避免与Python字符串定界符冲突
LQ = '\u201c'  # 左双引号 "
RQ = '\u201d'  # 右双引号 "


def set_font(run, size_pt, color=None, bold=False, name=None):
    run.font.size = Pt(size_pt)
    run.font.name = name or FONT_NAME_ASCII
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml('<w:rPr %s></w:rPr>' % nsdecls('w'))
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s w:eastAsia="%s"/>' % (nsdecls('w'), FONT_NAME))
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_paragraph_with_font(doc, text, size_pt, color=None, bold=False,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=0, space_after=6, line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    run = p.add_run(text)
    set_font(run, size_pt, color, bold)
    return p


def add_heading_styled(doc, text, level, size_pt, color=BLUE_TITLE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(24)
        pf.space_after = Pt(12)
    elif level == 2:
        pf.space_before = Pt(18)
        pf.space_after = Pt(8)
    else:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, size_pt, color, bold=True)
    return p


def add_body_text(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Cm(0.74)

    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        is_bold = (i % 2 == 1)
        set_font(run, 10.5, BLACK, bold=is_bold)
    return p


def add_image_with_caption(doc, image_filename, caption_text, width_cm=14):
    img_path = os.path.join(IMAGE_DIR, image_filename)
    if not os.path.exists(img_path):
        add_body_text(doc, '[图片缺失: %s]' % image_filename)
        return

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_img.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.0
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Cm(width_cm))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf_cap = p_cap.paragraph_format
    pf_cap.space_before = Pt(2)
    pf_cap.space_after = Pt(12)
    pf_cap.line_spacing = 1.0
    run_cap = p_cap.add_run(caption_text)
    set_font(run_cap, 9, GRAY_CAPTION, bold=False)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.2
    pf.left_indent = Cm(1)

    pPr = p._element.get_or_add_pPr()
    shd = parse_xml('<w:shd %s w:fill="F0F0F0" w:val="clear"/>' % nsdecls('w'))
    pPr.append(shd)

    run = p.add_run(code_text)
    set_font(run, 9.5, RGBColor(0x2D, 0x2D, 0x2D), name='Consolas')
    r = run._element
    rPr = r.find(qn('w:rPr'))
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is not None:
        rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, 10, BLUE_TITLE, bold=True)
        shading = parse_xml('<w:shd %s w:fill="D6E4F0" w:val="clear"/>' % nsdecls('w'))
        cell._element.get_or_add_tcPr().append(shading)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            set_font(run, 10, BLACK)

    return table


def build_document():
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ══════════════════════════════════════════════════════════
    #  封面页
    # ══════════════════════════════════════════════════════════

    for _ in range(6):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    add_paragraph_with_font(doc, '种子源模块控制器', 28, BLUE_TITLE, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_before=0, space_after=4, line_spacing=1.3)
    add_paragraph_with_font(doc, '— 软件设计文档 —', 28, BLUE_TITLE, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_before=0, space_after=24, line_spacing=1.3)

    add_paragraph_with_font(doc, '基于Qt 6.5框架的种子源模块上位机软件', 16, RGBColor(0x40, 0x40, 0x40),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_before=12, space_after=36, line_spacing=1.5)

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf_line = p_line.paragraph_format
    pf_line.space_before = Pt(12)
    pf_line.space_after = Pt(12)
    run_line = p_line.add_run('━' * 40)
    set_font(run_line, 10, RGBColor(0x1F, 0x4E, 0x79))

    add_paragraph_with_font(doc, '版本：v1.0.1', 14, RGBColor(0x40, 0x40, 0x40),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_before=24, space_after=8, line_spacing=1.5)
    add_paragraph_with_font(doc, '2026年6月', 14, RGBColor(0x40, 0x40, 0x40),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_before=8, space_after=0, line_spacing=1.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  第1章 设计总览
    # ══════════════════════════════════════════════════════════

    add_heading_styled(doc, '第1章 设计总览', 1, 18)

    add_body_text(doc,
        '本章从宏观视角介绍软件的整体设计思路，帮助读者在深入细节之前建立对系统全貌的认知。'
        '我们将从项目背景出发，阐述设计哲学，并说明关键技术选型的考量。')

    # 1.1
    add_heading_styled(doc, '1.1 项目背景与目标', 2, 14)

    add_body_text(doc,
        '种子源模块控制器是一款运行在Windows平台上的上位机软件，其核心目标是通过串口通信对种子源模块设备进行实时监控和精确控制。'
        '软件严格遵循《种子源模块通信协议V1.3》规范，在保证通信可靠性的同时，为用户提供直观的操作界面和丰富的数据可视化能力。', indent=True)

    add_body_text(doc,
        '从工程需求出发，本软件需要解决三个核心问题：第一，如何在与设备进行可靠串口通信的同时，保持用户界面的流畅响应；'
        '第二，如何将复杂的通信协议细节对用户屏蔽，提供简洁的操作抽象；'
        '第三，如何对设备运行数据进行有效的采集、存储和可视化展示。', indent=True)

    # 1.2
    add_heading_styled(doc, '1.2 设计哲学', 2, 14)

    add_body_text(doc,
        '本软件的设计遵循' + LQ + '分层解耦、线程隔离、信号驱动' + RQ + '三大原则。', indent=True)

    add_body_text(doc,
        '**分层解耦**意味着软件被划分为八个职责明确的层次，每一层只依赖其下层提供的接口，而不关心下层的实现细节。'
        '这种设计使得协议层的实现可以被替换（策略模式），通信层可以独立演化，而UI层完全不需要了解串口通信的细节。', indent=True)

    add_body_text(doc,
        '**线程隔离**是保证UI流畅响应的关键决策。我们将所有串口I/O操作放在独立的通信线程中执行，主线程只负责UI渲染和用户交互。'
        '两个线程之间通过Qt的信号槽机制（QueuedConnection）进行异步通信，从根本上避免了串口阻塞导致界面卡顿的问题。', indent=True)

    add_body_text(doc,
        '**信号驱动**是整个软件的运行模式。从用户点击按钮到设备状态更新，从数据接收到图表刷新，所有模块间的交互都通过Qt信号槽完成。'
        '这种松耦合的设计使得每个模块都可以独立开发和测试，也使得数据流的方向清晰可追踪。', indent=True)

    # 1.3
    add_heading_styled(doc, '1.3 技术选型', 2, 14)

    add_body_text(doc,
        '技术选型需要综合考虑开发效率、运行性能、维护成本和团队技术栈。下表列出了本项目各技术领域的选型及理由。')

    add_table(doc,
        ['技术领域', '选型', '选型理由'],
        [
            ['开发框架', 'Qt 6.5+', '成熟的跨平台框架，内置串口、数据库、图表支持'],
            ['编程语言', 'C++11', 'Qt原生语言，性能优异，适合实时系统'],
            ['编译器', 'MSVC 2019/2022', 'Windows平台最佳兼容性'],
            ['构建工具', 'CMake', '跨平台构建标准，灵活配置'],
            ['数据库', 'SQLite', '轻量级嵌入式数据库，无需额外部署'],
            ['图表库', 'QCustomPlot', '高性能实时绘图，支持缩放平移'],
        ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  第2章 系统架构
    # ══════════════════════════════════════════════════════════

    add_heading_styled(doc, '第2章 系统架构', 1, 18)

    add_body_text(doc,
        '系统架构是软件的骨架。本章将详细阐述分层架构的设计、多线程模型的实现以及数据流的走向，'
        '这三者共同构成了软件的运行基础。')

    # 2.1
    add_heading_styled(doc, '2.1 分层架构总览', 2, 14)

    add_image_with_caption(doc, 'system_architecture.png', '图2-1 系统分层架构图')

    add_body_text(doc,
        '如上图所示，软件采用八层分层架构。理解这个架构的关键在于把握每一层的职责边界和数据流向。', indent=True)

    add_body_text(doc,
        '最顶层的**UI层**是用户与软件交互的唯一入口。六个功能面板（连接、控制、状态、图表、报警、日志）各自承担明确的交互职责，'
        '它们不直接访问通信层或协议层，而是通过MainWindow作为中介来协调所有交互逻辑。'
        '这种设计确保了UI面板的独立性——任何一个面板的修改都不会影响其他面板或底层逻辑。', indent=True)

    add_body_text(doc,
        '**可视化层**封装了QCustomPlot图表控件，将实时数据渲染为直观的曲线图。它只关心数据本身，不关心数据从何而来。', indent=True)

    add_body_text(doc,
        '**数据模型层**是整个软件的数据枢纽。DeviceDataModel以观察者模式运作：当通信层接收到新的设备数据时，数据模型被更新，'
        '随后自动通知所有订阅者（UI面板、数据库等）。这种' + LQ + '推' + RQ + '模式的数据分发避免了UI层轮询数据的低效做法。', indent=True)

    add_body_text(doc,
        '**通信层**是架构中最复杂的部分。CommunicationWorker运行在独立线程中，管理着命令队列、串口收发、帧解析和超时重试。'
        '它通过三组信号（connectRequested/disconnectRequested/sendCommandRequested）接收主线程的请求，'
        '通过另外四组信号（connectionStateChanged/commandCompleted/logMessage/dataReceived）将结果反馈给主线程。', indent=True)

    add_body_text(doc,
        '**协议层**将《种子源模块通信协议V1.3》的具体实现封装在SeedSourceProtocolParser中。'
        '通过IProtocolParser接口（策略模式），协议实现可以被替换而不影响上层代码。'
        'RegisterManager作为单例管理所有寄存器定义，提供地址计算和值范围校验。', indent=True)

    add_body_text(doc,
        '底层的**报警层**、**数据库层**和**配置层**属于基础设施。它们以单例模式存在，为上层提供报警检测、数据持久化和配置管理服务。', indent=True)

    # 2.2
    add_heading_styled(doc, '2.2 多线程架构', 2, 14)

    add_image_with_caption(doc, 'thread_architecture.png', '图2-2 多线程架构图')

    add_body_text(doc,
        '多线程架构是本软件最核心的设计决策。上图清晰地展示了主线程和通信线程的职责划分与交互方式。', indent=True)

    add_body_text(doc,
        '**主线程**承载QApplication事件循环，负责所有UI渲染和用户交互。它持有MainWindow、六个UI面板、DeviceDataModel和轮询定时器'
        '（pollTimer，100ms周期）。当pollTimer触发时，MainWindow创建一个ReadStatusCommand并通过sendCommandRequested信号发送给通信线程'
        + LQ + '——这是设备数据周期性更新的驱动源。', indent=True)

    add_body_text(doc,
        '**通信线程**在run()方法中创建并管理QSerialPort和两个QTimer。pollTimer以10ms周期调用processCommandQueue()检查并执行待处理命令；'
        'connectionCheckTimer以5s周期检查串口连接状态。所有串口操作（open/close/read/write）都在这个线程中完成，主线程永远不会直接访问QSerialPort。', indent=True)

    add_body_text(doc,
        '两个线程之间的通信严格遵循Qt的信号槽机制，所有跨线程连接都使用Qt::QueuedConnection。'
        '这意味着信号发射后，槽函数在接收线程的事件循环中被调用，从而保证了线程安全。'
        '特别地，当通信线程解析出完整的协议帧后，它通过QMetaObject::invokeMethod(Qt::QueuedConnection)调用ICommand::onResponse()，'
        '确保onResponse()在ICommand所在的线程（通常是主线程）中执行，避免QTimer的跨线程访问问题。', indent=True)

    add_body_text(doc,
        '这种架构的一个关键细节是doDisconnectDevice()中断开所有pending命令的信号连接。'
        '如果不这样做，当命令对象被析构后，信号仍可能触发，导致use-after-free崩溃。', indent=True)

    # 2.3
    add_heading_styled(doc, '2.3 数据流', 2, 14)

    add_image_with_caption(doc, 'data_flow.png', '图2-3 数据流图')

    add_body_text(doc,
        '上图展示了从用户操作到UI更新的完整数据流。理解这条数据流是理解整个软件运作方式的关键。', indent=True)

    add_body_text(doc, '数据流可以分为三个阶段：', indent=True)

    add_body_text(doc,
        '**第一阶段：命令生成与发送（主线程→通信线程）**。用户在UI面板上的操作（如点击Start按钮）触发ControlPanel的信号，'
        'MainWindow的槽函数接收信号后，通过CommandFactory创建对应的命令对象。命令对象通过sendCommandRequested信号传递到通信线程，'
        '被加入命令队列。processCommandQueue()取出命令，调用command->execute()，execute()内部调用buildRequest()构建协议帧，'
        '然后通过sendRequest信号将帧数据传给sendNextRequest()写入串口。', indent=True)

    add_body_text(doc,
        '**第二阶段：响应接收与解析（通信线程）**。串口接收到数据后，onSerialDataReady()将数据追加到接收缓冲区。'
        'processReceivedData()在缓冲区中查找完整帧（findCompleteFrame()），找到后通过QMetaObject::invokeMethod调用command->onResponse()。'
        'onResponse()内部调用parseResponse()解析帧数据，提取电流、温度、功率等有效信息，然后发射completed信号。', indent=True)

    add_body_text(doc,
        '**第三阶段：数据分发与UI更新（通信线程→主线程）**。completed信号触发CommunicationWorker::onCommandCompleted()，'
        '该函数发射commandCompleted信号。MainWindow接收到此信号后，从result中提取数据并更新DeviceDataModel。'
        'DeviceDataModel发射dataUpdated信号，触发onDataUpdated()，进而更新StatusPanel、RealtimePlotPanel，并将数据保存到数据库。', indent=True)

    add_body_text(doc,
        '值得注意的是，整个数据流是单向的：命令从主线程流向通信线程，响应从通信线程流回主线程。'
        '这种单向数据流使得系统行为可预测、易调试。', indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  第3章 通信协议设计
    # ══════════════════════════════════════════════════════════

    add_heading_styled(doc, '第3章 通信协议设计', 1, 18)

    add_body_text(doc,
        '通信协议是软件与设备之间的' + LQ + '语言' + RQ + '。本章详细描述帧结构、命令体系和寄存器模型，'
        '这些内容构成了协议层实现的基础。')

    # 3.1
    add_heading_styled(doc, '3.1 帧结构', 2, 14)

    add_image_with_caption(doc, 'protocol_frame.png', '图3-1 通信协议帧结构图')

    add_body_text(doc,
        '通信协议帧遵循《种子源模块通信协议V1.3》规范，采用固定格式的帧结构。如上图所示，一帧数据由7个字段组成。', indent=True)

    add_body_text(doc,
        '帧头0xAA和帧尾0x55用于帧定界，接收端通过搜索帧头来定位帧的起始位置。地址码标识目标设备，默认为0x01。'
        '命令码定义了四种操作：0x01读寄存器、0x02写寄存器、0x03读状态、0x04设备控制。'
        '数据长度字段指示后续数据载荷的字节数，使得接收端可以确定帧的结束位置。'
        '校验和覆盖从帧头到数据载荷的所有字节，采用累加和取低8位的算法，用于检测传输错误。', indent=True)

    add_body_text(doc,
        '帧解析算法（findCompleteFrame）的设计考虑了实际通信中的各种异常情况：帧头前的垃圾数据会被丢弃；'
        '如果缓冲区超过4096字节仍无有效帧头，则清空缓冲区防止内存耗尽；如果帧尾不匹配，则跳过当前帧头继续搜索下一个。', indent=True)

    # 3.2
    add_heading_styled(doc, '3.2 命令体系', 2, 14)

    add_body_text(doc,
        '软件定义了四种命令类型，每种命令对应不同的数据载荷格式和响应解析逻辑。', indent=True)

    add_body_text(doc,
        '**读寄存器命令（0x01）**：发送2字节的基地址和偏移地址，接收4字节的大端序寄存器值。用于读取设备的各种参数。', indent=True)

    add_body_text(doc,
        '**写寄存器命令（0x02）**：发送2字节地址加4字节大端序值，接收确认响应。用于设置电流、温度、功率等参数。', indent=True)

    add_body_text(doc,
        '**读状态命令（0x03）**：无发送参数，接收12字节的设备状态数据（电流4字节+温度4字节+功率4字节，均为大端序）。'
        '这是轮询操作使用的命令，每100ms执行一次。', indent=True)

    add_body_text(doc,
        '**设备控制命令（0x04）**：发送1字节控制码（0x01启动/0x02停止/0x03重置/0x04校准），接收确认响应。', indent=True)

    # 3.3
    add_heading_styled(doc, '3.3 寄存器模型', 2, 14)

    add_body_text(doc,
        '设备定义了8个寄存器，每个寄存器由基地址和偏移地址唯一标识。地址计算公式为：', indent=True)

    add_code_block(doc, '实际地址 = 基地址 << 8 | 偏移地址')

    add_body_text(doc,
        'RegisterManager以单例模式管理所有寄存器定义，提供值范围校验和类型检查。'
        '这种集中式的寄存器管理确保了协议定义的一致性，避免了硬编码地址散落在各处带来的维护风险。', indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  第4章 核心模块设计
    # ══════════════════════════════════════════════════════════

    add_heading_styled(doc, '第4章 核心模块设计', 1, 18)

    add_body_text(doc,
        '本章深入四个核心模块的内部设计：命令系统、通信工作线程、数据模型和信号槽连接网络。'
        '这些模块是软件运行的中枢，理解它们的设计是理解整个软件的关键。')

    # 4.1
    add_heading_styled(doc, '4.1 命令系统', 2, 14)

    add_image_with_caption(doc, 'class_hierarchy.png', '图4-1 命令系统类继承关系图')

    add_body_text(doc,
        '命令系统是通信层与协议层之间的桥梁，采用命令模式设计。上图展示了类继承关系和关键接口。', indent=True)

    add_body_text(doc,
        'ICommand是所有命令的抽象基类，定义了三个核心接口：execute()负责构建请求帧并发送，'
        'onResponse()负责解析响应帧，onTimeout()处理超时情况。每个命令对象拥有唯一的原子递增ID（std::atomic<quint32>），'
        '用于在待处理命令映射中追踪。', indent=True)

    add_body_text(doc,
        '命令对象具有明确的生命周期状态机：Pending（待发送）→ Sending（构建请求中）→ WaitingResponse（等待响应）'
        '→ Completed/Error/Timeout（终态）。状态转换通过setState()方法触发，同时发射stateChanged信号。', indent=True)

    add_body_text(doc,
        '命令系统的线程安全设计值得特别说明。ICommand对象在主线程中创建（其QTimer也在主线程中），'
        '但execute()从通信线程调用。因此，QTimer的start()和stop()操作通过QMetaObject::invokeMethod(Qt::QueuedConnection)跨线程调用，'
        '确保QTimer始终在其所属线程中操作。同样，onResponse()也通过invokeMethod调用，'
        '使得parseResponse()在ICommand所在的线程中执行。', indent=True)

    add_body_text(doc,
        'CommandFactory以工厂模式封装了命令对象的创建逻辑，返回QSharedPointer<ICommand>智能指针，'
        '自动管理命令对象的生命周期。', indent=True)

    # 4.2
    add_heading_styled(doc, '4.2 通信工作线程', 2, 14)

    add_body_text(doc,
        'CommunicationWorker是整个软件中最复杂的模块，它承担了串口通信的全部职责。', indent=True)

    add_body_text(doc,
        '**线程生命周期**：startCommunication()启动线程，run()方法中创建QSerialPort和QTimer，进入exec()事件循环。'
        'stopCommunication()调用quit()退出事件循环，run()方法在exec()返回后清理QSerialPort和QTimer。'
        '析构函数调用stopCommunication()和wait()确保线程安全退出。', indent=True)

    add_body_text(doc,
        '**命令队列处理**：采用生产者-消费者模式。主线程通过sendCommandRequested信号将命令投递到队列，'
        '通信线程的processCommandQueue()以10ms周期检查队列。当通信状态为Idle且设备已连接时，取出队首命令执行。', indent=True)

    add_body_text(doc,
        '**帧解析**：接收缓冲区采用QMutex保护。findCompleteFrame()实现了鲁棒的帧搜索算法，'
        '处理了垃圾数据、帧头重复、帧尾不匹配等异常情况。', indent=True)

    add_body_text(doc,
        '**连接管理**：5秒定时器检查串口连接状态，如果检测到断开则发射connectionStateChanged信号通知主线程。', indent=True)

    # 4.3
    add_heading_styled(doc, '4.3 数据模型', 2, 14)

    add_body_text(doc,
        'DeviceDataModel采用观察者模式，是数据从通信层流向UI层的枢纽。', indent=True)

    add_body_text(doc,
        'updateData()方法的设计体现了对线程安全和性能的考量：先获取写锁更新数据，将数据添加到缓存，'
        '然后释放写锁，最后在无锁状态下发射信号。这种' + LQ + '先更新后通知' + RQ + '的模式避免了信号接收者在回调中访问到不一致的数据状态。', indent=True)

    add_body_text(doc,
        'DataCache内部类实现了环形缓冲区，默认容量10000条。当缓存满时，最旧的数据被自动丢弃。'
        'QReadWriteLock保护使得多个读取者可以并发访问缓存，而写入者独占访问。', indent=True)

    add_body_text(doc,
        'checkAlarms()方法实现了位域报警检测。6种报警类型分别对应报警字的6个位（0x01~0x20），'
        '当报警状态发生变化时发射alarmTriggered信号。', indent=True)

    # 4.4
    add_heading_styled(doc, '4.4 信号槽连接网络', 2, 14)

    add_image_with_caption(doc, 'signal_slot.png', '图4-2 信号槽连接网络图')

    add_body_text(doc,
        '上图展示了MainWindow为中心的信号槽连接网络。理解这张图就理解了整个软件的运行机制。', indent=True)

    add_body_text(doc,
        'MainWindow作为信号槽的中枢，连接了UI面板、通信线程、数据模型、数据库和报警管理器之间的所有交互。'
        '这种设计使得各模块之间没有直接依赖，所有交互都通过信号槽间接完成。', indent=True)

    add_body_text(doc, '信号流可以分为四条主线：', indent=True)

    add_body_text(doc,
        '**控制流**（绿色）：UI面板的用户操作信号 → MainWindow槽函数 → 创建命令 → 发送到通信线程。'
        '这是用户意图转化为设备操作的主通道。', indent=True)

    add_body_text(doc,
        '**通信流**（橙色）：通信线程的connectionStateChanged/commandCompleted/logMessage信号 → MainWindow槽函数 → 更新UI状态。'
        '这是设备状态反馈到界面的通道。', indent=True)

    add_body_text(doc,
        '**数据流**（蓝色）：commandCompleted → MainWindow更新DeviceDataModel → dataUpdated信号 → UI面板刷新和数据库保存。'
        '这是设备数据从接收到展示的通道。', indent=True)

    add_body_text(doc,
        '**报警流**（红色）：DeviceDataModel的alarmTriggered信号 → MainWindow → AlarmPanel显示和数据库保存。'
        '这是异常情况的处理通道。', indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  第5章 设计模式应用
    # ══════════════════════════════════════════════════════════

    add_heading_styled(doc, '第5章 设计模式应用', 1, 18)

    add_body_text(doc,
        '本软件系统地运用了六种设计模式，每种模式都解决了特定的设计问题。设计模式不是装饰，而是应对复杂度的利器。')

    add_body_text(doc,
        '**命令模式**是最核心的模式。ICommand将' + LQ + '对设备的操作' + RQ + '封装为对象，使得操作可以被排队、撤销、超时和追踪。'
        '如果没有命令模式，通信线程将需要大量的条件判断来区分不同类型的操作，代码将变得难以维护。', indent=True)

    add_body_text(doc,
        '**观察者模式**解决了数据分发的难题。DeviceDataModel作为被观察者，UI面板、数据库等作为观察者。'
        '当数据更新时，模型不需要知道谁在监听，只需要发射信号。这种松耦合使得新增数据消费者'
        '（如未来的远程监控模块）不需要修改模型代码。', indent=True)

    add_body_text(doc,
        '**策略模式**体现在IProtocolParser接口上。虽然目前只有SeedSourceProtocolParser一个实现，'
        '但如果将来需要支持新版本的通信协议，只需要实现新的解析器类即可，无需修改命令系统或通信线程的代码。', indent=True)

    add_body_text(doc,
        '**工厂模式**通过CommandFactory封装了命令对象的创建逻辑。调用者不需要知道具体创建哪种命令子类，'
        '只需要指定意图（读寄存器、写寄存器等），工厂方法负责选择正确的子类并构造对象。', indent=True)

    add_body_text(doc,
        '**单例模式**用于RegisterManager、DatabaseManager、AlarmManager和ConfigManager。'
        '这些管理器需要在整个应用生命周期中保持唯一实例，且需要被多个模块访问。单例模式配合互斥锁保证了线程安全。', indent=True)

    add_body_text(doc,
        '**生产者-消费者模式**体现在CommunicationWorker的命令队列中。主线程作为生产者投递命令，'
        '通信线程作为消费者执行命令。这种模式解耦了命令的生成速率和执行速率，使得主线程不会因为通信延迟而阻塞。', indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  第6章 线程安全设计
    # ══════════════════════════════════════════════════════════

    add_heading_styled(doc, '第6章 线程安全设计', 1, 18)

    add_body_text(doc,
        '线程安全是本软件设计中最需要谨慎对待的方面。我们采取了以下策略：')

    add_body_text(doc,
        '**原则一：对象归属线程**。每个QObject只在其创建线程中操作。QSerialPort、QTimer在通信线程中创建和销毁；'
        'ICommand的QTimer在主线程中创建，通过invokeMethod跨线程操作。', indent=True)

    add_body_text(doc,
        '**原则二：互斥锁保护共享数据**。CommunicationWorker使用四把互斥锁分别保护不同维度的共享数据：'
        'm_stateMutex保护连接状态和通信状态，m_queueMutex保护命令队列，m_pendingMutex保护待处理命令映射，'
        'm_bufferMutex保护接收缓冲区。这种细粒度锁策略减少了锁竞争，提高了并发性能。', indent=True)

    add_body_text(doc,
        '**原则三：跨线程调用使用invokeMethod**。当通信线程需要调用主线程对象的方法时（如ICommand::onResponse），'
        '使用QMetaObject::invokeMethod(Qt::QueuedConnection)，确保方法在目标线程的事件循环中执行。', indent=True)

    add_body_text(doc,
        '**原则四：断开信号防止悬空引用**。doDisconnectDevice()中断开所有pending命令的信号连接，'
        '防止命令对象析构后信号仍触发导致的use-after-free问题。', indent=True)

    add_body_text(doc,
        'ConfigManager使用QRecursiveMutex支持递归加锁（因为saveSettings()在setXxx()中被调用，而两者都需要加锁）。'
        'DeviceDataModel和AlarmManager使用QReadWriteLock，允许多个读取者并发访问。', indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  第7章 项目文件结构
    # ══════════════════════════════════════════════════════════

    add_heading_styled(doc, '第7章 项目文件结构', 1, 18)

    add_body_text(doc,
        '下表列出了项目所有源代码文件及其职责，帮助开发者快速定位代码位置。')

    add_table(doc,
        ['文件名', '职责说明'],
        [
            ['main.cpp', '应用程序入口，初始化并启动主窗口'],
            ['MainWindow.h/cpp', '主窗口，协调所有模块的交互中枢'],
            ['CommunicationWorker.h/cpp', '通信工作线程，管理串口收发和命令队列'],
            ['ICommand.h/cpp', '命令抽象基类，定义execute/onResponse/onTimeout接口'],
            ['ReadRegisterCommand.h/cpp', '读寄存器命令实现'],
            ['WriteRegisterCommand.h/cpp', '写寄存器命令实现'],
            ['ReadStatusCommand.h/cpp', '读状态命令实现'],
            ['DeviceControlCommand.h/cpp', '设备控制命令实现'],
            ['CommandFactory.h/cpp', '命令工厂，封装命令对象创建逻辑'],
            ['IProtocolParser.h', '协议解析器接口（策略模式）'],
            ['SeedSourceProtocolParser.h/cpp', '种子源协议V1.3解析器实现'],
            ['RegisterManager.h/cpp', '寄存器定义管理器（单例）'],
            ['DeviceDataModel.h/cpp', '设备数据模型，观察者模式数据枢纽'],
            ['ConnectionPanel.h/cpp', '连接管理UI面板'],
            ['ControlPanel.h/cpp', '设备控制UI面板'],
            ['StatusPanel.h/cpp', '设备状态显示UI面板'],
            ['RealtimePlotPanel.h/cpp', '实时图表UI面板'],
            ['AlarmPanel.h/cpp', '报警显示UI面板'],
            ['LogPanel.h/cpp', '日志显示UI面板'],
            ['DatabaseManager.h/cpp', 'SQLite数据库管理器（单例）'],
            ['AlarmManager.h/cpp', '报警检测管理器（单例）'],
            ['ConfigManager.h/cpp', '配置管理器（单例）'],
            ['CMakeLists.txt', 'CMake构建配置文件'],
        ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  第8章 版本历史
    # ══════════════════════════════════════════════════════════

    add_heading_styled(doc, '第8章 版本历史', 1, 18)

    add_body_text(doc, '以下是本软件的版本演进记录：')

    add_table(doc,
        ['版本号', '日期', '变更说明'],
        [
            ['v1.0.0', '2025-12-20', '初始版本发布'],
            ['v1.0.1', '2026-06-05', '重构通信线程架构，修复线程安全bug，集成windeployqt部署'],
        ])

    # ── 保存文档 ──
    doc.save(OUTPUT_PATH)
    print('文档已生成: %s' % OUTPUT_PATH)


if __name__ == '__main__':
    build_document()

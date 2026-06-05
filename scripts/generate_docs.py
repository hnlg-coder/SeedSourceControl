#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成「种子源模块控制器」设计文档与使用说明书 (v2.0)"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ============================================================
# 通用样式工具
# ============================================================

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, "F2F7FB")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_code_block(doc, code_text, language=""):
    """添加代码块（灰底等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # 灰底
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F5F5"/>')
    pPr.append(shd)
    return p

def add_ascii_diagram(doc, diagram):
    """添加ASCII艺术图"""
    return add_code_block(doc, diagram)

def add_tip(doc, text, tip_type="info"):
    """添加提示框"""
    colors = {"info": "E8F4FD", "warning": "FFF3CD", "danger": "F8D7DA"}
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    prefix = {"info": "ℹ️ ", "warning": "⚠️ ", "danger": "❌ "}
    run = p.add_run(prefix.get(tip_type, "") + text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{colors.get(tip_type, "E8F4FD")}"/>')
    pPr.append(shd)
    return p

# ============================================================
# Part 1: 软件设计文档
# ============================================================

def generate_design_doc():
    doc = Document()

    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- 样式 ---
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = '微软雅黑'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if i == 1:
            hs.font.size = Pt(18)
            hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif i == 2:
            hs.font.size = Pt(14)
            hs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            hs.font.size = Pt(11.5)
            hs.font.color.rgb = RGBColor(0x37, 0x84, 0xC4)

    # ========================
    # 封面
    # ========================
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("种子源模块控制器")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("软件设计文档")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 Qt 6.5 框架 · 多线程架构 · 通信协议 V1.3")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"版本：v2.0　|　{datetime.date.today().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ========================
    # 目录页（占位）
    # ========================
    doc.add_heading("目录", level=1)
    toc_items = [
        "第1章 设计总览",
        "  1.1 项目背景与目标",
        "  1.2 设计哲学",
        "  1.3 技术选型",
        "第2章 系统架构",
        "  2.1 分层架构总览",
        "  2.2 多线程架构",
        "  2.3 数据流",
        "第3章 通信协议设计",
        "  3.1 帧结构",
        "  3.2 命令体系",
        "  3.3 寄存器映射",
        "第4章 核心模块设计",
        "  4.1 命令系统 (ICommand)",
        "  4.2 通信工作线程 (CommunicationWorker)",
        "  4.3 数据模型 (DeviceDataModel)",
        "  4.4 协议解析器 (SeedSourceProtocolParser)",
        "  4.5 可视化引擎 (EnhancedChartWidget)",
        "  4.6 统计计算器 (StatisticsCalculator)",
        "第5章 UI 模块设计",
        "  5.1 主窗口与页签导航",
        "  5.2 各功能页签",
        "第6章 线程安全设计",
        "第7章 设计模式应用",
        "第8章 项目文件结构",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ========================
    # 第1章 设计总览
    # ========================
    doc.add_heading("第1章 设计总览", level=1)

    doc.add_heading("1.1 项目背景与目标", level=2)
    doc.add_paragraph(
        "种子源模块控制器是一款运行在 Windows 平台上的上位机软件，通过串口（UART）与种子源"
        "模块设备进行双向通信，实现设备参数配置、实时状态监控、数据可视化分析及历史数据管理。"
    )
    doc.add_paragraph(
        "软件严格遵循《种子源模块通信协议 V1.3》规范，采用多线程架构确保通信可靠性与界面"
        "流畅度。核心目标包括：可靠串口通信、直观的多页签操作界面、高性能实时曲线渲染、"
        "滑动窗口统计分析、以及 SQLite 数据持久化。"
    )

    doc.add_heading("1.2 设计哲学", level=2)
    doc.add_paragraph("本软件设计遵循三大原则：")

    p = doc.add_paragraph()
    run = p.add_run("分层解耦：")
    run.bold = True
    p.add_run("软件分为 UI 层、可视化层、数据模型层、通信层、协议层、数据库层、配置层，"
              "每层仅依赖其下层接口，不关心下层实现细节。协议层的实现可通过 IProtocolParser "
              "策略接口替换，通信层可独立演化。")

    p = doc.add_paragraph()
    run = p.add_run("线程隔离：")
    run.bold = True
    p.add_run("所有阻塞式串口 I/O 操作运行在独立通信线程中，主线程专注于 UI 渲染与用户交互。"
              "两线程间通过 Qt QueuedConnection 信号槽异步通信，GUI 在任何情况下不会因"
              "串口阻塞而卡顿。图表渲染同样使用独立 RenderThread，进一步保障主线程流畅度。")

    p = doc.add_paragraph()
    run = p.add_run("信号驱动：")
    run.bold = True
    p.add_run("模块间交互全部通过 Qt 信号槽机制完成。从用户点击按钮到设备状态更新，"
              "从数据接收到图表刷新，所有事件流清晰可追踪、松耦合、可独立测试。")

    doc.add_heading("1.3 技术选型", level=2)
    add_table(doc,
        ["技术领域", "选型", "选型理由"],
        [
            ["开发框架", "Qt 6.5 (Widgets)", "成熟的跨平台 C++ GUI 框架，原生 Windows 渲染"],
            ["C++ 标准", "C++11", "与项目复杂度匹配，利用 std::atomic 等特性"],
            ["编译器", "MSVC 2019/2022", "Windows 平台标准工具链"],
            ["构建系统", "CMake 3.16+", "跨平台构建，支持 AUTOMOC"],
            ["串口通信", "QSerialPort", "Qt 官方串口模块，跨平台，API 简洁"],
            ["数据存储", "SQLite (QSqlDatabase)", "嵌入式数据库，零配置，适合单机应用"],
            ["图表渲染", "自研 EnhancedChartWidget", "独立渲染线程 + QImage 离屏渲染，避免 QCustomPlot 许可证问题"],
            ["文档生成", "python-docx", "程序化生成 .docx，与代码同步维护"],
        ],
        col_widths=[3, 4, 9.5]
    )

    doc.add_page_break()

    # ========================
    # 第2章 系统架构
    # ========================
    doc.add_heading("第2章 系统架构", level=1)

    doc.add_heading("2.1 分层架构总览", level=2)
    doc.add_paragraph(
        "软件采用八层分层架构，各层职责明确、单向依赖。下图展示了完整的层次关系"
        "及模块归属。"
    )

    add_ascii_diagram(doc,
r"""┌─────────────────────────────────────────────────────┐
│                    UI 层 (Qt Widgets)                 │
│  ┌──────────┐ ┌──────────────┐ ┌──────────────────┐  │
│  │ 7个页签   │ │ ConnectionPanel│ │    LogPanel     │  │
│  │ (总览/电流 │ │ DashboardWidget│ │ (嵌入总览页签)   │  │
│  │ /温度/功率 │ │ CurrentCtrlWidget│ │                │  │
│  │ /配置/报警 │ │ TemperatureCtrl │ │ DataTablePanel │  │
│  │ /统计)    │ │ MonitorWidget  │ │                │  │
│  └──────────┘ └──────────────┘ └──────────────────┘  │
│              MainWindow (信号槽中枢)                    │
├─────────────────────────────────────────────────────┤
│              可视化层 (EnhancedChartWidget)            │
│         RenderThread + QImage 离屏渲染                 │
│         StatisticsCalculator 滑动窗口统计               │
├─────────────────────────────────────────────────────┤
│          数据模型层 (DeviceDataModel)                   │
│     观察者模式 | StatusBits/AlertBits/ConfigBits      │
│     DataCache 环形缓冲区 (10000条)                     │
├─────────────────────────────────────────────────────┤
│                  通信层 (CommunicationWorker : QThread) │
│     命令队列 (生产者-消费者) | 帧解析 | 超时重试        │
│     5×QMutex 细粒度锁 | QueuedConnection 跨线程通信    │
├─────────────────────────────────────────────────────┤
│              协议层 (IProtocolParser 策略接口)          │
│     SeedSourceProtocolParser: 帧构造/解析/校验         │
│     符合《种子源模块通信协议 V1.3》                    │
├─────────────────────────────────────────────────────┤
│       基础设施层: 数据库 | 报警 | 配置 | 日志           │
│     DatabaseManager 单例  SQLite 持久化                │
└─────────────────────────────────────────────────────┘""")

    doc.add_paragraph()
    doc.add_paragraph(
        "架构的核心思想是单向依赖：上层可调用下层接口，下层通过信号槽向上层通知事件，"
        "但下层绝不直接依赖上层的具体实现。这种设计使得每一层可独立编译、测试和替换。"
    )

    doc.add_heading("2.2 多线程架构", level=2)
    doc.add_paragraph(
        "多线程架构是本软件最核心的设计决策，直接决定了系统的响应性能和稳定性。"
        "以下是线程分配模型："
    )

    add_ascii_diagram(doc,
r"""┌────────────────── 主线程 (GUI Thread) ──────────────────┐
│  QApplication::exec()                                     │
│  ┌────────────┐    signals          ┌──────────────────┐  │
│  │ MainWindow │───connectRequested──▶│                  │  │
│  │            │───disconnectRequested▶│ CommunicationWorker│
│  │  100ms     │───sendCommandRequested▶│   (子线程)       │  │
│  │  pollTimer │                     │                  │  │
│  │            │◀──connectionStateChanged──│  QSerialPort    │  │
│  │            │◀──logMessage──────────│  QTimer×2       │  │
│  │            │◀──commandCompleted────│  命令队列        │  │
│  │            │◀──dataReceived───────│  帧缓冲区        │  │
│  └────────────┘                     └──────────────────┘  │
│       │                                                    │
│       ▼                                                    │
│  ┌─────────────────────────────────────────────┐          │
│  │ DeviceDataModel ──▶ dataUpdated ──▶ 7个页签    │          │
│  └─────────────────────────────────────────────┘          │
└────────────────────────────────────────────────────────────┘

┌────────────── RenderThread ──────────────┐
│  QWaitCondition 驱动                      │
│  QImage 离屏绘制 → paintEvent 贴图         │
│  多系列曲线 + 坐标轴 + 图例 + 统计面板      │
└────────────────────────────────────────────┘""")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("线程安全关键措施：")
    run.bold = True

    bullets = [
        "所有跨线程信号连接使用 Qt::QueuedConnection，确保槽函数在目标线程事件循环中执行。",
        "CommunicationWorker 使用 5 把 QMutex 分别保护：连接状态、命令队列、待处理映射、接收缓冲区、串口配置。",
        "ICommand::m_state 使用 std::atomic<CommandState> + compare_exchange_strong 防止双重发射 completed 信号。",
        "EnhancedChartWidget::RenderThread 的数据通过独立 QMutex 保护，渲染参数使用 std::atomic<int> 原子变量。",
        "MainWindow::m_connected 和 m_running 使用 std::atomic<bool>，消除主线程与子线程间的数据竞争。",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading("2.3 数据流", level=2)
    doc.add_paragraph("从用户操作到 UI 刷新的完整数据流可分为三个阶段：")

    add_ascii_diagram(doc,
r"""阶段1: 命令生成 ── 主线程
  用户操作 → UI Widget 信号 → MainWindow 槽函数
    → CommandFactory::createXxx() → sendCommandRequested 信号

阶段2: 命令执行 ── 通信线程
  doSendCommand() → command->execute()
    → buildRequest() [协议层构造帧]
    → 串口 write() → 等待响应 → 串口 read()
    → findCompleteFrame() [帧解析/校验]
    → command->onResponse() [CAS 原子状态切换]

阶段3: 数据分发 ── 主线程
  commandCompleted 信号 → QVariantMap 解析
    → DeviceDataModel::updateData() [写锁]
    → dataUpdated 信号 [无锁发射]
    → 7 个页签 Widget 更新 + DatabaseManager 持久化""")

    doc.add_page_break()

    # ========================
    # 第3章 通信协议设计
    # ========================
    doc.add_heading("第3章 通信协议设计", level=1)

    doc.add_heading("3.1 帧结构", level=2)
    doc.add_paragraph("通信帧遵循《种子源模块通信协议 V1.3》，帧结构如下：")

    add_ascii_diagram(doc,
r"""偏移    0        1        2        3      4..4+L-1   4+L     5+L
     ┌───────┬───────┬───────┬───────┬───────┬───────┬───────┐
     │ 0xAA  │ ADDR  │  CMD  │ LEN   │  DATA │  CRC  │ 0x55  │
     │ 帧头   │ 地址码 │ 命令码 │ 数据长 │ 数据载荷│ 校验和 │ 帧尾   │
     └───────┴───────┴───────┴───────┴───────┴───────┴───────┘
     固定值   1字节   1字节   1字节    L字节   1字节   固定值

ADDR = (基地址 << 4) | (偏移地址 & 0x0F)
CRC  = 从帧头到数据载荷所有字节累加和取低 8 位""")

    doc.add_paragraph()
    add_table(doc,
        ["字段", "长度", "说明"],
        [
            ["帧头", "1 字节", "固定 0xAA，用于帧同步"],
            ["ADDR", "1 字节", "地址 = (baseAddr<<4) | offset"],
            ["CMD", "1 字节", "命令码，定义见 3.2"],
            ["LEN", "1 字节", "数据载荷字节数"],
            ["DATA", "LEN 字节", "变长数据载荷，大端序"],
            ["CRC", "1 字节", "累加校验和 (帧头→DATA 累加取低8位)"],
            ["帧尾", "1 字节", "固定 0x55"],
        ],
        col_widths=[2, 2.5, 12]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("帧解析算法 (findCompleteFrame) 容错设计：")
    run.bold = True
    items = [
        "帧头前的垃圾数据自动丢弃",
        "缓冲区超过 4096 字节仍未找到有效帧头 → 清空缓冲区防内存溢出",
        "校验和不匹配 → 跳过当前帧头，从 headerIndex+1 继续搜索",
        "帧尾不匹配 → 继续等待数据或跳过",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_heading("3.2 命令体系", level=2)

    add_table(doc,
        ["命令码", "命令名", "数据载荷（发送）", "响应载荷", "用途"],
        [
            ["0xF0", "读寄存器", "READ_LEN(1B)", "寄存器值(4B, 大端)", "读取设备寄存器值"],
            ["0x0F", "写寄存器", "WRITE_LEN(1B) + VALUE(NB)", "确认响应", "设置设备参数"],
            ["0x03", "读状态", "无", "STATUS+ALERT+CUR+TEMP+POWER(各4B)", "批量读取设备状态(轮询)"],
            ["0x04", "设备控制", "ACT_CODE(1B): 01启动/02停止/03复位/04校准", "确认响应", "控制设备运行状态"],
            ["0x00", "心跳", "无", "心跳响应", "检测设备在线"],
            ["0xAE", "告警上报", "—", "ALERT(4B)", "设备主动上报告警"],
        ],
        col_widths=[1.5, 2, 4, 3.5, 5.5]
    )

    doc.add_heading("3.3 寄存器映射", level=2)
    doc.add_paragraph("设备寄存器按基地址 + 偏移地址编址，ADDR = (baseAddr << 4) | offset：")

    add_table(doc,
        ["基地址", "偏移", "ADDR", "寄存器名", "位宽", "说明"],
        [
            ["0x01", "0x00", "0x10", "DEVINFO", "32-bit", "设备信息(name/ver/serial)"],
            ["0x01", "0x01", "0x11", "SYSTEM", "8-bit", "系统控制(sleep/reset/iap)"],
            ["0x02", "0x00", "0x20", "STATUS", "8-bit", "运行状态(idle/cur/temp)"],
            ["0x03", "0x00", "0x30", "CONFIG", "64-bit", "配置(使能位+阈值+斜率+波特率)"],
            ["0x04", "0x00", "0x40", "TEMP", "16-bit", "温度设定值 & 实时值"],
            ["0x05", "0x00", "0x50", "CUR", "32-bit", "电流设定值 & 实时值"],
            ["0x06", "0x00", "0x60", "MONITOR", "32-bit", "功率监测(激光/恒流/系统)"],
            ["0x07", "0x00", "0x70", "ALERT", "8-bit", "告警标志(8个告警位)"],
        ],
        col_widths=[1.8, 1.5, 1.5, 2.5, 1.8, 7.5]
    )

    doc.add_page_break()

    # ========================
    # 第4章 核心模块设计
    # ========================
    doc.add_heading("第4章 核心模块设计", level=1)

    doc.add_heading("4.1 命令系统 (ICommand)", level=2)
    doc.add_paragraph(
        "命令系统采用命令模式（Command Pattern），将设备操作封装为可排队、可追踪、可超时的对象。"
        "类继承关系如下："
    )

    add_ascii_diagram(doc,
r"""                    ┌──────────────┐
                    │   QObject    │
                    └──────┬───────┘
                           │
                    ┌──────▼───────┐
                    │   ICommand   │  ← 抽象基类
                    │──────────────│
                    │ + execute()  │  buildRequest() → sendRequest 信号
                    │ + onResponse()│  parseResponse() → CAS 状态切换
                    │ + onTimeout() │  CAS 防护 → completed 信号
                    │──────────────│
                    │ m_state:     │  std::atomic<CommandState>
                    │ m_result:    │  QVariant + QMutex 保护
                    │ s_nextId:    │  std::atomic<quint32> 唯一ID
                    └──────┬───────┘
           ┌──────┬───────┼───────┬──────┐
           │      │       │       │      │
    ┌──────▼─┐ ┌──▼───┐ ┌─▼────┐ ┌▼──────┐
    │ReadReg │ │WriteReg│ │ReadSt│ │CtrlDev│
    │Command │ │Command │ │Command│ │Command│
    └────────┘ └───────┘ └──────┘ └───────┘""")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("线程安全状态机（核心改进）：")
    run.bold = True

    add_ascii_diagram(doc,
r"""Pending → Sending → WaitingResponse → Processing → Completed
                  ↓                    ↓
               Error               Timeout

关键: onResponse/onTimeout 使用 compare_exchange_strong 原子操作
      确保 completed 信号只发射一次，彻底消除双重发射竞态""")

    doc.add_paragraph()
    doc.add_paragraph(
        "ICommand 对象由 CommandFactory 工厂类创建，返回 QSharedPointer<ICommand>。"
        "命令生命周期由智能指针自动管理，pending 命令存储在 CommunicationWorker 的 "
        "QMap<quint32, QSharedPointer<ICommand>> 中，通过唯一 ID 追踪。"
    )

    doc.add_heading("4.2 通信工作线程 (CommunicationWorker)", level=2)
    doc.add_paragraph(
        "CommunicationWorker 继承 QThread，是串口通信的唯一执行者。"
        "其核心职责包括：命令队列管理、串口收发、帧解析、超时处理、连接状态监控。"
    )

    add_table(doc,
        ["组件", "说明"],
        [
            ["命令队列 (QQueue)", "主线程投递命令 → 子线程取出执行（生产者-消费者模式）"],
            ["待处理映射 (QMap)", "cmdId → 命令对象，用于响应匹配"],
            ["接收缓冲区 (QByteArray)", "串口数据累积，findCompleteFrame() 提取完整帧"],
            ["m_pollTimer (10ms)", "processCommandQueue() 周期检查并执行命令"],
            ["m_connectionCheckTimer (5s)", "检测串口连接状态，异常断开时通知主线程"],
            ["5×QMutex", "m_queueMutex/m_pendingMutex/m_bufferMutex/m_stateMutex/m_configMutex"],
        ],
        col_widths=[4.5, 12]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("连接建立流程：")
    run.bold = True

    add_ascii_diagram(doc,
r"""MainWindow                  CommunicationWorker         Device
   │                               │                      │
   │───connectRequested 信号────────▶│                      │
   │                               │───open(COMx)─────────▶│
   │                               │◀───handle 返回─────────│
   │◀──connectionStateChanged(true)─│                      │
   │                               │  start pollTimer(10ms) │
   │                               │  start connCheck(5s)  │""")

    doc.add_heading("4.3 数据模型 (DeviceDataModel)", level=2)
    doc.add_paragraph(
        "DeviceDataModel 是数据中枢，采用观察者模式。当通信层接收到设备数据后，"
        "模型被更新，随后自动通知所有订阅者（UI 页签、数据库）。"
    )

    add_table(doc,
        ["结构体", "对应寄存器", "位域定义"],
        [
            ["StatusBits", "STATUS[02h]", "idle(bit0) | curStatus(bit1-2) | tempStatus(bit3-4)"],
            ["AlertBits", "ALERT[07h]", "ccCtrl/ccPd/ccAdc/ccDac/tcCtrl/tcAdc/tcDac/sysPwr"],
            ["ConfigBits", "CONFIG[03h]", "tcEn/ccEn/aeEn/powerSv/curSv/tempSv(bit0-5) + curTh/curSlope/baudRate"],
            ["DevInfo", "DEVINFO[01h]", "name/verS/verH/serial/curMax/tempMin/tempMax"],
            ["RealTimeData", "综合", "13+ 字段，含兼容字段 current/temperature/power/statusRaw/alarmRaw"],
        ],
        col_widths=[2.5, 2.8, 11.3]
    )

    doc.add_paragraph()
    add_code_block(doc,
r"""// 数据更新与信号发射的线程安全模式
void DeviceDataModel::updateData(...) {
    QWriteLocker locker(&m_dataLock);     // 1. 获取写锁，更新数据
    m_currentData.current = current;
    // ... 更新所有字段 ...
    m_dataCache->add(m_currentData);      // 2. 添加到环形缓存
    locker.unlock();                       // 3. 释放写锁
    checkAlarms();                         // 4. 无锁状态检查告警
    emit dataUpdated(m_currentData);       // 5. 无锁状态发射信号
}""")

    doc.add_heading("4.4 协议解析器 (SeedSourceProtocolParser)", level=2)
    doc.add_paragraph(
        "实现 IProtocolParser 策略接口，封装《种子源模块通信协议 V1.3》的全部帧构造"
        "与解析逻辑。通过策略模式，未来可替换为其他协议版本而不修改上层代码。"
    )

    add_code_block(doc,
r"""// ADDR 编码: (基地址 << 4) | (偏移 & 0x0F)
// 例: baseAddr=0x05(CUR寄存器), offset=0x00
//     → ADDR = (0x05 << 4) | 0x00 = 0x50

// 写寄存器帧构造:
// DATA = WRITE_LEN(1B) + VALUE(大端序, N字节)
// CMD = 0x0F""")

    doc.add_heading("4.5 可视化引擎 (EnhancedChartWidget)", level=2)
    doc.add_paragraph(
        "EnhancedChartWidget 是自研高性能图表控件，使用独立 RenderThread + QImage "
        "离屏渲染方案，避免主线程阻塞。支持多系列曲线叠加、鼠标交互缩放、统计信息叠加显示。"
    )

    add_ascii_diagram(doc,
r"""┌─── EnhancedChartWidget ──────────────────────────┐
│  ┌─ RenderThread ──────────────────────────────┐  │
│  │  QWaitCondition + QMutex                     │  │
│  │  ┌──────────────────────────────────┐       │  │
│  │  │ renderToImage():                  │       │  │
│  │  │  QPainter on QImage              │       │  │
│  │  │  ├─ 网格线 + 坐标轴刻度          │       │  │
│  │  │  ├─ 多条数据曲线 (QVector<Point>)│       │  │
│  │  │  ├─ 图例 (颜色+名称)             │       │  │
│  │  │  └─ 统计面板 (均值/方差/最大/最小)│       │  │
│  │  └──────────────────────────────────┘       │  │
│  │  frameReady(QImage) → 主线程 paintEvent     │  │
│  └──────────────────────────────────────────────┘  │
│  鼠标交互: 滚轮缩放Y轴 / 拖拽平移(预留)            │
└────────────────────────────────────────────────────┘""")

    doc.add_heading("4.6 统计计算器 (StatisticsCalculator)", level=2)
    doc.add_paragraph(
        "基于滑动窗口的实时统计计算器，为每条数据曲线提供统计分析能力。"
        "默认窗口大小 1000 个数据点，线程安全（QReadWriteLock 保护）。"
    )

    add_table(doc,
        ["统计指标", "计算方式"],
        [
            ["平均值 (mean)", "窗口内所有值的算术平均"],
            ["方差 (variance)", "Σ(xi - mean)² / n"],
            ["标准差 (stdDev)", "sqrt(方差)"],
            ["最大值 (max)", "窗口内最大值"],
            ["最小值 (min)", "窗口内最小值"],
            ["最新值 (last)", "最近一个数据点"],
            ["计数 (count)", "当前窗口内数据点数"],
        ],
        col_widths=[3.5, 13]
    )

    doc.add_page_break()

    # ========================
    # 第5章 UI 模块设计
    # ========================
    doc.add_heading("第5章 UI 模块设计", level=1)

    doc.add_heading("5.1 主窗口与页签导航", level=2)
    doc.add_paragraph(
        "MainWindow 继承 QMainWindow，使用 QTabWidget 作为核心导航控件，包含 7 个功能页签。"
        "窗口支持自由缩放（最小 800×600），布局自动适配。"
    )

    add_ascii_diagram(doc,
r"""┌─── 种子源模块控制器 - v2.0 ──────────────────────┐
│ 文件(&F)  帮助(&H)                                  │
├────────────────────────────────────────────────────┤
│ Port:[COM1▼] Baud:[115200▼] ...[Refresh][Connect]  │
│ Status: Disconnected                               │
├────────────────────────────────────────────────────┤
│ [总览] [电流控制] [温度控制] [功率监测] [设备配置]   │
│ [报警管理] [数据统计]                                │
├────────────────────────────────────────────────────┤
│                                                    │
│        当前选中页签的内容区域                         │
│        (含图表/控件/日志等)                          │
│                                                    │
└────────────────────────────────────────────────────┘""")

    add_table(doc,
        ["页签", "核心组件", "功能说明"],
        [
            ["总览", "DashboardWidget + LogPanel", "设备状态一览、启停控制、运行日志"],
            ["电流控制", "CurrentControlWidget", "目标电流设定、实际电流显示、电流趋势曲线"],
            ["温度控制", "TemperatureControlWidget", "目标温度设定、实际温度显示、温度趋势曲线"],
            ["功率监测", "MonitorWidget", "激光/恒流/系统功率双曲线图"],
            ["设备配置", "ConfigWidget", "CONFIG 位配置、DEVINFO 查看、SYSTEM 控制"],
            ["报警管理", "AlertWidget", "8 路告警指示灯、告警历史记录表"],
            ["数据统计", "DataTablePanel", "实时数据表格、CSV 导出（最多 1000 行）"],
        ],
        col_widths=[2, 4.5, 10]
    )

    doc.add_page_break()

    # ========================
    # 第6章 线程安全设计
    # ========================
    doc.add_heading("第6章 线程安全设计", level=1)

    doc.add_paragraph("线程安全是系统稳定性的基石。当前版本已通过以下措施全面加固：")

    add_table(doc,
        ["原则", "实现", "覆盖范围"],
        [
            ["对象线程归属", "QObject 子类仅在创建线程操作", "全局"],
            ["std::atomic 状态变量", "m_state(CommandState)/m_connected/m_running/m_renderRequested", "ICommand, MainWindow, RenderThread"],
            ["CAS 防双重发射", "compare_exchange_strong 守卫 onResponse/onTimeout", "ICommand"],
            ["QMutex 保护共享数据", "5 把锁保护通信层，QReadWriteLock 保护数据模型", "CommunicationWorker, DeviceDataModel"],
            ["QMutex 保护结果数据", "m_result/m_errorString 加锁读写", "ICommand"],
            ["QueuedConnection 跨线程", "所有跨线程信号槽连接", "MainWindow ↔ CommWorker"],
            ["析构顺序保证", "stopCommunication → wait → delete", "MainWindow::~MainWindow()"],
            ["Meyers 单例自动析构", "static 局部变量确保构造/析构线程安全", "DatabaseManager"],
        ],
        col_widths=[3.5, 6.5, 6.5]
    )

    doc.add_page_break()

    # ========================
    # 第7章 设计模式应用
    # ========================
    doc.add_heading("第7章 设计模式应用", level=1)

    add_table(doc,
        ["设计模式", "应用位置", "解决的问题"],
        [
            ["命令模式", "ICommand 体系", "将设备操作封装为对象，支持排队、超时、追踪"],
            ["观察者模式", "DeviceDataModel → dataUpdated 信号", "数据变更自动通知所有订阅者，解耦数据生产与消费"],
            ["策略模式", "IProtocolParser 接口", "协议实现可替换，不影响上层代码"],
            ["工厂模式", "CommandFactory", "封装命令对象创建，调用者无需关心具体子类"],
            ["单例模式", "DatabaseManager", "全局唯一数据库连接，线程安全构造与析构"],
            ["生产者-消费者", "命令队列 (QQueue)", "解耦命令投递速率与执行速率"],
        ],
        col_widths=[2.5, 5, 9]
    )

    doc.add_page_break()

    # ========================
    # 第8章 项目文件结构
    # ========================
    doc.add_heading("第8章 项目文件结构", level=1)

    add_ascii_diagram(doc,
r"""SeedSourceControl/
├── CMakeLists.txt                    # CMake 构建配置
├── main.cpp                          # 应用入口
├── README.md
├── resources/
│   └── default.qss                   # 全局 Qt 样式表
├── src/
│   ├── mainwindow.h/.cpp             # 主窗口
│   ├── communication/
│   │   ├── command.h/.cpp            # ICommand 命令体系
│   │   └── communicationworker.h/.cpp # 通信工作线程
│   ├── protocol/
│   │   ├── protocolparser.h/.cpp     # 协议V1.3 编解码
│   ├── model/
│   │   ├── devicedatamodel.h/.cpp    # 数据模型（观察者）
│   ├── ui/
│   │   ├── connectionpanel.h/.cpp    # 串口连接面板
│   │   ├── dashboardwidget.h/.cpp    # 总览页签
│   │   ├── currentcontrolwidget.h/.cpp # 电流控制页签
│   │   ├── temperaturecontrolwidget.h/.cpp # 温度控制页签
│   │   ├── monitorwidget.h/.cpp      # 功率监测页签
│   │   ├── configwidget.h/.cpp       # 设备配置页签
│   │   ├── alertwidget.h/.cpp        # 报警管理页签
│   │   ├── datatablepanel.h/.cpp     # 数据统计页签
│   │   └── logpanel.h/.cpp           # 日志面板
│   ├── visualization/
│   │   ├── enhancedchartwidget.h/.cpp # 图表渲染引擎
│   │   └── statisticscalculator.h/.cpp # 统计计算器
│   └── database/
│       └── databasemanager.h/.cpp    # SQLite 数据库管理
├── build/                            # 构建输出
└── docs/                             # 设计文档 & 使用说明书""")

    add_table(doc,
        ["目录", "职责"],
        [
            ["src/communication/", "通信层：命令系统 + 通信线程"],
            ["src/protocol/", "协议层：V1.3 帧编解码"],
            ["src/model/", "数据模型层：状态/告警/配置位域 + 数据缓存"],
            ["src/ui/", "UI 层：7 个页签 + 连接面板 + 日志面板"],
            ["src/visualization/", "可视化层：图表渲染 + 统计计算"],
            ["src/database/", "数据库层：SQLite 持久化"],
        ],
        col_widths=[4, 12.5]
    )

    # --- 保存 ---
    output_path = os.path.join(OUTPUT_DIR, "种子源模块控制器_软件设计文档_v2.0.docx")
    doc.save(output_path)
    print(f"[OK] 设计文档已保存: {output_path}")
    return output_path


# ============================================================
# Part 2: 使用说明书
# ============================================================

def generate_user_manual():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = '微软雅黑'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if i == 1:
            hs.font.size = Pt(18)
            hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif i == 2:
            hs.font.size = Pt(14)
            hs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            hs.font.size = Pt(11.5)
            hs.font.color.rgb = RGBColor(0x37, 0x84, 0xC4)

    # --- 封面 ---
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("种子源模块控制器")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("使 用 说 明 书")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"版本：v2.0　|　{datetime.date.today().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # ========================
    # 1. 软件简介
    # ========================
    doc.add_heading("1. 软件简介", level=1)

    add_table(doc,
        ["项目", "说明"],
        [
            ["软件名称", "种子源模块控制器 (SeedSourceControl)"],
            ["版本", "v2.0"],
            ["运行环境", "Windows 10/11（64位）"],
            ["软件用途", "通过串口与种子源模块设备通信，实现参数配置、实时监控、数据可视化与历史数据管理"],
            ["通信协议", "《种子源模块通信协议 V1.3》"],
        ],
        col_widths=[3, 13.5]
    )

    doc.add_paragraph()
    doc.add_paragraph("主要功能：")
    items = [
        "设备连接通信：串口参数配置，连接/断开管理，连接状态实时检测",
        "参数配置：电流、温度、功率目标值设定，设备配置寄存器读写",
        "实时数据监控：7 个功能页签，全面展示设备运行状态",
        "数据可视化：独立渲染线程驱动的实时曲线图，支持多系列叠加显示",
        "报警管理：8 路告警指示灯 + 告警历史记录，位域精确诊断",
        "数据管理：实时数据表格（1000 行）、CSV 导出、SQLite 自动持久化",
        "统计分析：滑动窗口统计（平均值、方差、最大值、最小值）",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    # ========================
    # 2. 安装与启动
    # ========================
    doc.add_heading("2. 安装与启动", level=1)

    doc.add_heading("2.1 系统要求", level=2)
    items = [
        "操作系统：Windows 10/11 64位（Windows 7 理论上兼容但未经测试）",
        "硬件：至少一个可用串口（物理 COM 口或 USB 转串口）",
        "软件依赖：已集成所有 Qt 运行时依赖，无需单独安装",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_heading("2.2 启动方式", level=2)
    doc.add_paragraph("双击 build/deploy/Debug/SeedSourceControl.exe 即可启动程序。")
    doc.add_paragraph("程序以 Windows 标准 GUI 窗口模式运行。")

    doc.add_heading("2.3 目录说明", level=2)
    add_table(doc,
        ["文件/目录", "说明"],
        [
            ["SeedSourceControl.exe", "主程序"],
            ["platforms/", "Qt 平台插件 (qwindows.dll)"],
            ["sqldrivers/", "SQLite 数据库驱动 (qsqlite.dll)"],
            ["styles/", "Qt 样式插件"],
            ["imageformats/", "图像格式插件"],
        ],
        col_widths=[4, 12.5]
    )

    doc.add_page_break()

    # ========================
    # 3. 界面说明
    # ========================
    doc.add_heading("3. 界面说明", level=1)

    doc.add_heading("3.1 主界面布局", level=2)
    doc.add_paragraph("主界面从上到下依次为：菜单栏 → 串口连接面板 → 页签导航栏 → 内容区域。")

    add_ascii_diagram(doc,
r"""┌─── 种子源模块控制器 - v2.0 ──────────────────────┐
│ 文件(&F)  帮助(&H)                    ← 菜单栏      │
├────────────────────────────────────────────────────┤
│ Port:[▼] Baud:[▼] Data:[▼] Parity:[▼] Stop:[▼]     │
│ Flow:[▼] [Refresh] [Connect]  Status:Disconnected  │
│ ← 串口连接面板                                       │
├────────────────────────────────────────────────────┤
│ [总览] [电流控制] [温度控制] [功率监测] [设备配置]   │
│ [报警管理] [数据统计]              ← 页签导航        │
├────────────────────────────────────────────────────┤
│                                                    │
│            当前选中页签的内容区域                      │
│    (含状态显示 / 参数设定 / 实时曲线 / 日志等)        │
│                                                    │
└────────────────────────────────────────────────────┘""")

    doc.add_heading("3.2 菜单栏", level=2)
    add_table(doc,
        ["菜单", "功能"],
        [
            ["文件(&F) → 退出(&X)", "退出程序（快捷键 Ctrl+Q）"],
            ["帮助(&H) → 关于(&A)", "显示软件版本和版权信息"],
        ],
        col_widths=[5, 11.5]
    )

    doc.add_heading("3.3 串口连接面板", level=2)

    add_table(doc,
        ["控件", "说明"],
        [
            ["Port", "串口选择下拉框，自动检测系统可用 COM 端口，支持手动刷新"],
            ["Baud Rate", "波特率：9600/14400/19200/38400/57600/115200/230400/460800"],
            ["Data Bits", "数据位：5/6/7/8（默认 8）"],
            ["Parity", "校验位：None/Even/Odd/Mark/Space（默认 None）"],
            ["Stop Bits", "停止位：1/1.5/2（默认 1）"],
            ["Flow Control", "流控制：None/Hardware/Software（默认 None）"],
            ["Refresh", "刷新可用串口列表"],
            ["Connect/Disconnect", "连接/断开设备（连接后按钮文字切换）"],
            ["Status 标签", "连接状态显示：绿色 Connected / 红色 Disconnected"],
        ],
        col_widths=[3.5, 13]
    )

    add_tip(doc, "连接后所有配置控件将被禁用，需断开后才能修改串口参数。")

    doc.add_heading("3.4 总览页签 (Dashboard)", level=2)
    doc.add_paragraph("设备运行状态总览，包含以下区域：")

    items = [
        "状态指示灯：设备状态（IDLE/RUN/ERR）、电流状态、温度状态",
        "数值显示：设定电流、实际电流、设定温度、实际温度、功率、报警状态",
        "快捷控制：[启动]（绿色）/[停止]（红色）按钮",
        "运行日志：位于页签底部，以深色背景展示。包含 Info/Warning/Error 三级日志，"
        "显示通信 TX/RX 十六进制数据和操作记录。支持 [清除] 和 [保存] 操作",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_heading("3.5 电流控制页签", level=2)
    doc.add_paragraph("电流参数设定与实时监测：")
    items = [
        "目标电流设定：QDoubleSpinBox，范围 0–1000.00 mA，步进 0.01",
        "[应用] 按钮：将设定值通过写寄存器命令发送到设备",
        "实际电流显示：实时刷新的大字数值标签",
        "电流趋势曲线：EnhancedChartWidget 驱动，支持鼠标滚轮缩放 Y 轴",
        "统计分析：曲线右侧显示均值/方差/最大/最小值",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_heading("3.6 温度控制页签", level=2)
    doc.add_paragraph("温度参数设定与实时监测：")
    items = [
        "目标温度设定：QDoubleSpinBox，范围 -40.00–125.00 °C，步进 0.01",
        "[应用] 按钮：将设定值通过写寄存器命令发送到设备",
        "实际温度显示：实时刷新的大字数值标签（>80°C 红色，>60°C 橙色）",
        "温度趋势曲线：同 EnhancedChartWidget",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_heading("3.7 功率监测页签", level=2)
    doc.add_paragraph("功率分布监测，包含双曲线图：")
    items = [
        "电流分配曲线：PD 电流、TEC 电流两个系列叠加显示",
        "功率分配曲线：激光功率、恒流功率、系统总功率三个系列叠加显示",
        "数值标签：curPd/curTec/pwrLas/pwrCc/pwrSys 五个独立指标",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_heading("3.8 设备配置页签", level=2)
    doc.add_paragraph("CONFIG 寄存器配置、DEVINFO 查看、SYSTEM 控制：")


    add_table(doc,
        ["区域", "控件", "功能"],
        [
            ["CONFIG", "6 个 CheckBox", "TC_EN/CC_EN/AE_EN/POWER_SV/CUR_SV/TEMP_SV 使能位"],
            ["CONFIG", "电流阈值/电流斜率", "QDoubleSpinBox 设定"],
            ["CONFIG", "波特率", "QComboBox 选择 9600-460800"],
            ["CONFIG", "[应用]" , "读取所有控件值，组合为 ConfigBits，发送写寄存器命令"],
            ["DEVINFO", "7 个 Label", "显示设备名称/硬件版本/软件版本/序列号/最大电流/温度范围"],
            ["DEVINFO", "[读取设备信息]", "发送读 DEVINFO 寄存器命令（功能预留）"],
            ["SYSTEM", "[休眠]/[软复位]/[固件升级]", "发送系统控制命令（功能预留）"],
        ],
        col_widths=[2, 4, 10.5]
    )

    doc.add_heading("3.9 报警管理页签", level=2)
    doc.add_paragraph("8 路告警指示灯 + 历史记录：")
    items = [
        "指示灯（8 组）：ccCtrl/ccPd/ccAdc/ccDac/tcCtrl/tcAdc/tcDac/sysPwr",
        "正常状态：绿色指示灯 + 文字 \"正常\"",
        "告警状态：红色指示灯 + 文字 \"告警\" + 告警描述",
        "告警历史表：时间 | 告警码 | 告警描述，记录每次告警触发事件",
        "[清除历史] 按钮：清除告警历史记录",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_heading("3.10 数据统计页签", level=2)
    doc.add_paragraph("实时数据表格与导出：")
    items = [
        "表格列：时间(yyyy-MM-dd HH:mm:ss.zzz) | 电流(mA) | 温度(°C) | 功率(mW) | 状态 | 报警",
        "最大行数：1000 行（超出后自动删除最旧行）",
        "[清除] 按钮：清空表格",
        "[导出CSV] 按钮：导出为 CSV 文件（UTF-8 编码），可用 Excel 打开",
        "导出文件名格式：data_yyyyMMdd_HHmmss.csv",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_page_break()

    # ========================
    # 4. 操作流程
    # ========================
    doc.add_heading("4. 操作流程", level=1)

    doc.add_heading("4.1 连接设备", level=2)
    steps = [
        "① 在串口面板的 Port 下拉框选择正确的 COM 端口。",
        "② 确认波特率等参数与设备一致（默认 115200-8-N-1）。",
        "③ 点击 [Connect] 按钮。",
        "④ 观察 Status 标签变为绿色 \"Connected\"。",
        "⑤ 日志面板显示 \"Device connected: COMx\"。",
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading("4.2 启动设备", level=2)
    steps = [
        "① 确认设备已连接。",
        "② 在总览页签点击 [启动]（绿色）按钮。",
        "③ 日志面板显示设备启动信息。",
        "④ 实时图表开始更新数据曲线。",
        "⑤ 各页签的状态数值开始动态刷新。",
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading("4.3 设置参数", level=2)
    steps = [
        "① 切换到电流控制页签（或温度控制页签）。",
        "② 在目标值输入框中输入期望的电流/温度值。",
        "③ 点击 [应用] 按钮发送写寄存器命令。",
        "④ 日志面板显示 \"设置目标电流/温度: xxx\"。",
        "⑤ 观察实际值标签和曲线的变化。",
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading("4.4 监控报警", level=2)
    steps = [
        "① 切换到报警管理页签，查看 8 路指示灯状态。",
        "② 当设备异常时，对应指示灯变红，自动记录到告警历史表。",
        "③ 日志面板同步记录告警详情（Error 级别）。",
        "④ 总览页签的 Alarm 状态同步更新。",
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading("4.5 停止和断开", level=2)
    steps = [
        "① 在总览页签点击 [停止]（红色）按钮。",
        "② 点击串口面板的 [Disconnect] 按钮。",
        "③ Status 标签变为红色 \"Disconnected\"。",
        "④ 所有配置控件恢复可编辑状态。",
        "⑤ 数据已自动保存到 SQLite 数据库。",
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_page_break()

    # ========================
    # 5. 数据管理
    # ========================
    doc.add_heading("5. 数据管理", level=1)

    doc.add_heading("5.1 数据存储", level=2)
    items = [
        "数据库类型：SQLite",
        "数据库位置：用户 AppData 目录",
        "data 表字段：id | timestamp | current | temperature | power | status | alarm",
        "alarms 表字段：id | timestamp | code | message",
        "自动保存：每次 dataUpdated 信号触发时自动写入",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_heading("5.2 数据导出", level=2)
    items = [
        "格式：CSV (UTF-8 with BOM)",
        "内容：时间 | 电流(mA) | 温度(°C) | 功率(mW) | 状态 | 报警",
        "操作：数据统计页签 → [导出CSV] 按钮",
        "排序：按时间戳升序",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_heading("5.3 内存缓存", level=2)
    items = [
        "DeviceDataModel 维护 10000 条数据的环形缓冲区",
        "EnhancedChartWidget 每条曲线维护 10000 个数据点",
        "StatisticsCalculator 滑动窗口默认 1000 个样本",
    ]
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

    doc.add_page_break()

    # ========================
    # 6. 常见问题
    # ========================
    doc.add_heading("6. 常见问题", level=1)

    faqs = [
        ("Q1: 无法找到串口",
         "检查设备是否已连接到电脑；检查 USB 转串口驱动是否已安装；"
         "点击 [Refresh] 按钮重新扫描；在设备管理器中确认 COM 端口正常识别。"),
        ("Q2: 连接失败",
         "确认串口参数与设备一致（默认 115200-8-N-1）；确认串口未被其他程序占用；"
         "检查设备是否已上电；查看日志面板的错误信息。"),
        ("Q3: 数据不更新",
         "确认设备已连接且已启动（在总览页签点击 [启动]）；"
         "检查轮询定时器是否正常（查看日志面板 TX/RX 数据）；确认波特率设置正确。"),
        ("Q4: 参数设置无效",
         "确认设备已连接且正在运行；参数变更需点击 [应用] 按钮发送到设备；"
         "检查参数值是否在有效范围内。"),
        ("Q5: 窗口太小内容显示不全",
         "窗口支持自由拖拽缩放（最小 800×600），拖拽窗口边缘即可调整大小；"
         "所有控件和图表会自动适配新尺寸。"),
        ("Q6: 如何导出数据",
         "切换到数据统计页签，点击 [导出CSV] 按钮；"
         "生成的 CSV 文件可用 Excel 直接打开。"),
        ("Q7: 图表上如何查看具体数值",
         "电流控制/温度控制/功率监测页签的图表支持鼠标滚轮缩放 Y 轴；"
         "每个图表的统计面板实时显示均值、方差、最大值、最小值。"),
    ]
    for title, answer in faqs:
        doc.add_heading(title, level=3)
        doc.add_paragraph(answer)

    doc.add_page_break()

    # ========================
    # 7. 技术参数
    # ========================
    doc.add_heading("7. 技术参数", level=1)

    add_table(doc,
        ["参数", "值"],
        [
            ["通信接口", "RS-232 / UART (QSerialPort)"],
            ["默认波特率", "115200 bps"],
            ["轮询周期", "200 ms"],
            ["命令超时", "2000 ms"],
            ["连接检测周期", "5000 ms"],
            ["数据缓存容量", "10000 条（环形缓冲区）"],
            ["表格最大行数", "1000 行"],
            ["统计窗口大小", "1000 样本"],
            ["电流设定范围", "0 – 1000.00 mA"],
            ["温度设定范围", "-40.00 – 125.00 °C"],
            ["功率监测范围", "0 – 10000.00 mW"],
            ["数据库类型", "SQLite 3"],
            ["CSV 编码", "UTF-8 with BOM"],
            ["Qt 版本", "6.5.3"],
            ["编译器", "MSVC 2019/2022"],
            ["C++ 标准", "C++11"],
        ],
        col_widths=[4, 12.5]
    )

    # --- 保存 ---
    output_path = os.path.join(OUTPUT_DIR, "种子源模块控制器_使用说明书_v2.0.docx")
    doc.save(output_path)
    print(f"[OK] 使用说明书已保存: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_design_doc()
    generate_user_manual()
